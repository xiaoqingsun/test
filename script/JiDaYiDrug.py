#!python
import glob
import sys
import imp
imp.reload(sys)
import re
import os
import time
import webbrowser
import xlrd
import argparse
import traceback
from collections import defaultdict
from docx import Document
from docx.shared import Inches,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def argInit():
    parser = argparse.ArgumentParser(
        description="report_create program.")
    parser.add_argument(
        '-dir', help='FWgxy_drug.pdf.txt', required=True)
    parser.add_argument(
        '-cfg', help='file cfg.txt', required=True)
    parser.add_argument(
        '-out', help='project name', default=os.getcwd())
    parser.add_argument(
        '-sample', help='sampleid', required=True)
    argv = vars(parser.parse_args())
    global drugfile,cfgFile,outDir,sampleid
    drugfile = argv['dir'].strip()
    cfgFile = argv['cfg'].strip()
    outDir = os.path.abspath(argv['out'].strip()) + "/"
    print (outDir)
    sampleid=argv['sample'].strip()

def get_genotype(drugfile,sampleid):
    filter_li={}
    if os.path.isfile(drugfile) == False:
        sys.stderr.write('%s is not a file' % drugfile)
        sys.exit(1)
    try:
        with open(drugfile, 'rU') as file_handle:
            for line in file_handle:
                line_list=[x.strip() for x in line.strip().split('\t')]
                key=line_list[2]+':'+line_list[8]
                filter_li[key]=line_list[9]+'\t'+line_list[10]+'\t'+line_list[25]
    except Exception as e:
        traceback.print_exc(e)
    return filter_li



def get_sampleInf(sampleid,cfgFile):
    sample_inf_list=[]
    cover_Dic = defaultdict(list)
    head=['sampleID','name','sex','age','data1','data2','doctor','sampletype']
    time2str=time.strftime('%Y-%m-%d',time.localtime(time.time()))
    if os.path.isfile(cfgFile) == False:
        print('%s is not a file' % cfgFile)
        return cover_Dic
    try:
        with open(cfgFile, 'rU') as file_handle:
            for line in file_handle:
                line_list=[x.strip() for x in line.strip().split('\t')]
                if line_list[0]==sampleid:
                    line_list[5]=re.sub(r'\(.*?\)','',line_list[5])
                    (age,date,sex,doc,samtyp)=get_HTB(sampleid) #wph changed 190423 修改性别不匹配的问题
                    if sex not in line_list[5]:
                        line_list[5] = sex
                    print (date)
                    print  (time2str)
                    sample_inf_list=[line_list[0],line_list[4],line_list[5],age,date,time2str,doc,samtyp]  #需要送样日期，暂时用出报告日期代替
                    break
        cover_Dic = {k:v for k,v in zip(head,sample_inf_list)}
    except Exception as e:
        traceback.print_exc(e)
    return cover_Dic



def get_HTB(sampleid):
    #files=glob.glob(r'/PUBLIC/pipline/database/sheet_hutong/*心康互通表*.xls*')
    files=glob.glob(r'/PUBLIC/pipline/database/sheet_hutong/2019心康送样信息单*.xls*')
    max_num=0
    age=date=''
    sex=''
    doc=''
    stype=''
    age_num = '' #wph changed 181106
    date_num = ''
    for i in files:
        hutong=i.split('/')[-1]
        #st=re.match(r'心康互通表（([0-9]*)）.xlsx',hutong)
        st=re.match(r'2019心康送样信息单-([0-9]*).xlsx',hutong)
        if int(st.group(1))> max_num:
            max_num=int(st.group(1))
    HTB=r'/PUBLIC/pipline/database/sheet_hutong/2019心康送样信息单-%s.xlsx' %max_num
    try:
        data = xlrd.open_workbook(HTB)
        table = data.sheet_by_name('下单表')
        nrows = table.nrows
        for rownum in range(1,nrows):
            if table.row(rownum)[2].value == sampleid:
                age=table.row(rownum)[8].value
                sex=table.row(rownum)[7].value
                doc=table.row(rownum)[10].value
                stype=table.row(rownum)[11].value
                ctype = table.cell(rownum, 1).ctype
                cell =table.cell_value(rownum, 1)
                date
                if ctype ==3 :
                    date = datetime(*xldate_as_tuple(cell, 0))
                    date = date.strftime('%Y/%m/%d')
                else:
                    date=table.row(rownum)[1].value
        date_num=date.replace("/","-")
        if isinstance(age, float) == True: #wph changed 181106
            age_num=str(int(age)) + "岁"
        elif not age:
            age_num="不详"
        else:
            age_num=age
        if age_num=='' or date_num=='':
            sys.stderr.write(u"ID不在互通表 或 互通表中没有年龄/送样日期信息")
            exit
    except Exception as e:
        traceback.print_exc(e)
    print (doc)
    print (stype)
    return age_num,date_num,sex,doc,stype

def main():
    argInit()
    context={}
    RSsite = get_genotype(drugfile,sampleid)
    sampleInf_Dic = get_sampleInf(sampleid,cfgFile)
    for key in sampleInf_Dic:
        context[key]=sampleInf_Dic[key]
    document = Document('/PUBLIC/pipline/script/Report/XKFW/V3_180824/tempt/吉大一用药报告-添加位点-模板.docx')
    table1=document.tables[0]
    table2=document.tables[1]
    table2.style.font.size=Pt(9) 
    table3=document.tables[2]
 
    table1.cell(0,0).text=u'姓名：'+context['name']
    table1.cell(0,1).text=u'性别：'+context['sex']
    table1.cell(0,2).text=u'年龄：'+context['age']
    table1.cell(0,3).text=u'样本类型：'+context['sampletype']
    table1.cell(2,3).text=u'样本号：'+context['sampleID']
    ###
    index=0
    for row in table2.rows:
        index=index+1
        if index==1:
            continue
        key=row.cells[1].text+':'+row.cells[3].text
        con=RSsite[key].split('\t')
        run = row.cells[4].paragraphs[0].add_run(con[0])
        run.font.size = Pt(9)
        run = row.cells[5].paragraphs[0].add_run(con[2])
        run.font.size = Pt(9)
        if con[1] =='-':
            con[1] = '正常'
        run = row.cells[6].paragraphs[0].add_run(con[1])
        run.font.size = Pt(9)
    ####   
    table3.cell(0,0).text=u'接收时间：'+context['data1']
    table3.cell(0,2).text=u'报告时间：'+context['data2']

    document.save(outDir+sampleid+'-吉大一用药报告.docx')


if __name__ == '__main__':
    main()
