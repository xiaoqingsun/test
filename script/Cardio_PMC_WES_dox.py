#-*- coding:utf-8 -*-
import os
import sys
import shutil
import logging
import argparse
import traceback
import time
import decimal
import xlrd
import os
import glob
import sys
import re
import codecs 
import lxml
from string import Template
from datetime import datetime
from xlrd import xldate_as_tuple
from collections import defaultdict
from docx import Document
from docx.shared import Mm, Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm,Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docxtpl import DocxTemplate, InlineImage, RichText

tpl=DocxTemplate('/PUBLIC/pipline/script/Report/XKFW/V3_180824/tempt/南方医院报告模板-确认版-20200102-muban-2.docx')

def argInit():
    parser = argparse.ArgumentParser(
        description="report_create program.")
    parser.add_argument(
        '-dir', help='file CS.filter.txt', required=True)
    parser.add_argument(
        '-cfg', help='file cfg.txt', required=True)
    parser.add_argument(
        '-feature', help='feature', required=True)
    parser.add_argument(
        '-out', help='project name', default=os.getcwd())
    parser.add_argument(
        '-qc', help='file Stat_QC.xls', required=True)
    parser.add_argument(
        '-sample', help='sampleid', required=True)
    parser.add_argument(
        '-genelist', help='genelist', required=True)
    parser.add_argument(
        '-module', help='module', required=True)
    parser.add_argument(
        '-yinyang', help='阴阳性', required=True)
    argv = vars(parser.parse_args())
    global filterFile,outDir,sampleid,module,yinyang_file,cfgFile,QCfile,feature_file,GeneListfile
    filterFile = argv['dir'].strip()
    module = argv['module'].strip()
    outDir = os.path.abspath(argv['out'].strip())
    cfgFile = argv['cfg'].strip()
    QCfile = argv['qc'].strip()
    GeneListfile = argv['genelist'].strip()
    sampleid=argv['sample'].strip()
    feature_file = argv['feature'].strip()
    yinyang_file=argv['yinyang'].strip()

def get_filter(filterFile,sampleid):
    filter_li=[]
    filter_fu=[]
    filter_bz=[]
    id_li=[]
    if os.path.isfile(filterFile) == False:
        sys.stderr.write('%s is not a file' % filterFile)
        sys.exit(1)
    try:
        with open(filterFile, 'rU') as file_handle:
            for line in file_handle:
                if re.match('^ID',line.strip()):continue
                if re.match('^####',line.strip()):continue
                line_list=[x.strip() for x in line.strip().split('\t')]
                if re.match('^###',line.strip()):
                    newsam=re.sub('###','',line_list[0])
                    if  newsam==sampleid:
                        line_1=re.sub('%','%',line)
                        line_2=re.sub('_','_',line_1)
                        line_3=re.sub('~','-',line_2) ### panqi 20181213
                        filter_bz.append(line_3.strip())
                elif re.match('^##',line.strip()): 
                    newsam=re.sub('##','',line_list[0])
                    if  newsam==sampleid:
                        line_1=re.sub('%','%',line)
                        line_2=re.sub('_','_',line_1)
                        line_3=re.sub('~','-',line_2) ### panqi 20181213
                        filter_fu.append(line_3.strip())
                else:
                    id_li.append(line_list[0])
                    if ''.join(line_list[1:6]) == '-----':continue
                    if line_list[0]==sampleid:
                        line_1=re.sub('%','%',line)
                        line_2=re.sub('_','_',line_1)
                        line_3=re.sub('~','-',line_2) ### panqi 20181213
                        filter_li.append(line_3.strip())
                    else:
                        pass
        if sampleid not in id_li:
            sys.stderr.write('sample %s is not in %s'%(sampleid,filterFile))
            sys.exit(1)
    except Exception as e:
        traceback.print_exc(e)
    return filter_li,filter_fu,filter_bz


def prepare_allDic(files,filesfu,filesbz):
    disease_Dic = defaultdict(list)
    site_Dic = defaultdict(list)
    referencearry= defaultdict(list)
    site_li = []
    gene_strand={}
    #disease_li= {}

    het_hom ={'het':u'杂合','hom':u'纯合','hem':u'半合子'} #wph add hem 181212
    #data_Dic=read_CSdata()   疾病名，疾病亚型，遗传模式，检出率，疾病类别 ，A类文献，B类文献
    #-f79致病性结论  80 疾病描述 81 基因描述 82 位点描述 83 参考文献 84 疾病名称 85 遗传模式 86 基因名称 87 CDS 88 PEP 89 PDF_NMID (GXY 92 PDF_NMID)
    NMID_index='-'
    if module == 'GXY':
        NMID_index=90
    elif module == 'IDT' or module == 'CS':
        NMID_index=88
    else:
        sys.stderr.write(u'请检查输入模块')
        sys.exit(1)
    
    
    try:
        for line in files:
            line_list=[x.strip() for x in line.strip().split('\t')]
            if len(line_list) < 88:
                sys.stderr.write('filter文件列数不足，请检查')
                sys.exit(1)
            ##获取disease_Dic
            level_value=get_level_value(line_list[78])
            detail=[line_list[83],line_list[84],line_list[79],level_value] #疾病名称、遗传模式、疾病描述、致病性等级对应的数字
            disease_Dic[line_list[85]].append(detail)  #key为基因,以基因为单位存储疾病,用作第一个疾病表格
               
            ##获取site_Dic
            site=line_list[1]+":"+line_list[2]+":"+line_list[3]+":"+line_list[4]+":"+line_list[5]+":"+line_list[83]+":"+line_list[84] #chr:start:end:ref:alt
            if site not in site_li:    #相同基因（不同疾病）有多个相同位点，都进行存储
                site2=line_list[1]+":"+line_list[2]+":"+line_list[3]+":"+line_list[4]+":"+line_list[5]  #chr:start:end:ref:alt
                chr_site=get_chr(line_list[1],line_list[2],line_list[3])
                chr_site2=line_list[1]+'-'+line_list[2]
                print (chr_site)
                if line_list[NMID_index]=='-':
                    exon='-'
                else:
                    ncbi_tran=line_list[NMID_index]   #changed by sunxq
                    exons=ncbi_tran.split(':')
                    exon=exons[0]+' '+exons[1]
                mute_type=get_mut_type(line_list[9])  
                #level_value=get_level_value(line_list[78])
                site_detail=[line_list[86],line_list[87],chr_site,exon,het_hom[line_list[31].lower()],mute_type,line_list[78],line_list[82],line_list[80],line_list[81],level_value,line_list[83],line_list[84],line_list[79],site2,line_list[20]] 
                             #nue,           acid,      chr_site,exon,het/hom/,                      mute_type,zhibingxing,    literature,  gene_des,     site_des,    level_value，疾病名称、    遗传模式、     疾病描述、  site，以位点为点为进行存储，,用于主表和位点解析的表格
                site_Dic[line_list[85]].append(site_detail)
                referencearry[site2].append([line_list[78],line_list[83],line_list[84],line_list[85],line_list[82]])
                gene_strand[chr_site2]=line_list[42]
                site_li.append(site)
                
    except Exception as e:
            traceback.print_exc(e)
    sort_disease = sort_Dic(disease_Dic)
    sort_site= sort_siteDic(site_Dic)
    

    disease_Dic_fubiao = defaultdict(list)
    site_Dic_fubiao = defaultdict(list)
    site_li_fubiao = []
    disease_li_fubiao = []
    try:
        for line in filesfu:
            line_list=[x.strip() for x in line.strip().split('\t')]
            site=line_list[1]+":"+line_list[2]+":"+line_list[3]+":"+line_list[4]+":"+line_list[5]+":"+line_list[83]+":"+line_list[84]
            site2=line_list[1]+":"+line_list[2]+":"+line_list[3]+":"+line_list[4]+":"+line_list[5] 
            chr_site2=line_list[1]+'-'+line_list[2]
            if site not in site_li_fubiao:
                if line_list[NMID_index]=='-':
                    exon='-'
                else:
                    ncbi_tran=line_list[NMID_index]   #changed by sunxq
                    exons=ncbi_tran.split(':')
                    exon=exons[0]+' '+exons[1]
                mute_type=get_mut_type(line_list[9])   
                level_value=get_level_value(line_list[78]) 
                chr_site=get_chr(line_list[1],line_list[2],line_list[3])
                varsite=chr_site+' '+line_list[86]+' '+line_list[87]
                print (varsite)
                site_detail=[varsite,exon,het_hom[line_list[31].lower()],line_list[20],line_list[77],line_list[78],level_value,line_list[83],line_list[84],site2]
                site_Dic_fubiao[line_list[85]].append(site_detail)
                site_li_fubiao.append(site)
                referencearry[site2].append([line_list[78],line_list[83],line_list[84],line_list[85],line_list[82]])
                gene_strand[chr_site2]=line_list[42]
               
            else:
                 pass
    except Exception as e:
            traceback.print_exc(e)
    Dic_bz = defaultdict(list)
    try:
        for line in filesbz:
             line_list=[x.strip() for x in line.strip().split('\t')]
             genekey=re.sub(' ','',line_list[85])
             Dic_bz[genekey].append(line_list[83]) 

    except Exception as e:
            traceback.print_exc(e)
    sort_disease_fubiao=disease_Dic_fubiao
    sort_site_fubiao = sort_siteDic_fubiao(site_Dic_fubiao)
    return sort_disease,sort_site,sort_disease_fubiao,sort_site_fubiao,Dic_bz,referencearry,gene_strand

def sort_Dic(Dic):
    sorted_Dic={}
    for key in Dic:
        sorted_Dic[key]=sorted(Dic[key],key=lambda x:x[3])
    return sorted_Dic

def sort_siteDic(Dic):
    sorted_Dic={}
    for key in Dic:
        sorted_Dic[key]=sorted(Dic[key],key=lambda x:x[10])
    return sorted_Dic

def sort_siteDic_fubiao(Dic):
    sorted_Dic={}
    for key in Dic:
        sorted_Dic[key]=sorted(Dic[key],key=lambda x:x[8])
    return sorted_Dic


def get_chr(Chr,start,end):
    site=''
    if start=='.' or start=='-' or Chr=='.' or Chr=='-':
        sys.stderr.write(u'染色体号或位点起始为空')
        sys.exit(1)
    if end=='.' or end=='-' or start == end:
        site='Chr'+Chr+':'+start
    else:
        site='Chr'+Chr+':'+start+' - '+end
    return site


def get_level_value(level):
    level_value=0
    if level=='致病': level_value=1
    elif level=='可能致病': level_value=2
    elif '临床意义未明1级' in re.sub(' ','',level): level_value=3
    elif '临床意义未明2级' in re.sub(' ','',level): level_value=4
    elif '临床意义未明3级' in re.sub(' ','',level): level_value=5
    elif '临床意义未明4级' in re.sub(' ','',level): level_value=6
    elif '临床意义未明5级' in re.sub(' ','',level): level_value=7
    elif level=='可能良性': level_value=8
    elif level=='良性': level_value=9
    else:
        sys.stderr.write(u'请检查致病等级')
        sys.exit(1)
    return level_value

def get_mut_type(mute):
    mute_Dic={'frameshift deletion':'移码缺失变异',
    'frameshift insertion':'移码插入变异',
    'frameshift substitution':'移码替换变异',
    'nonframeshift deletion':'非移码缺失变异',
    'nonframeshift insertion':'非移码插入变异',
    'nonframeshift substitution':'非移码替换变异',
    'stopgain':'无义变异',
    'stoploss':'终止缺失变异',
    'synonymous SNV':'同义变异',
    'nonsynonymous SNV':'错义变异',
    'splicing SNV':'剪切区单核苷酸变异',
    'splicing INDEL':'剪切区插入缺失变异',
    'intronic':'内含子变异',
    'SNV':'单核苷酸变异',
    '.':'.',
    }
    if mute in mute_Dic:
        mute_des=mute_Dic[mute]
    else:
        sys.stderr.write(u'请检查filter文件第9列是否为%s'%mute_Dic.keys())
        sys.exit(1)
    return mute_des

def get_zongshu(site_Dic):
    cover_Dic={}
    hom_gene,hom_level,het_gene,het_level = tongji_hethom(site_Dic)
    hom_gene_set=[x for x in set(hom_gene)]
    het_gene_set=[x for x in set(het_gene)]
    if len(het_gene)==0:    ##全为纯合突变
        if len(hom_gene)==1:
            zongshu='本次受检样本中共检出 1 个纯合变异，存在于%s基因上，'%(hom_gene[0])
        elif len(set(hom_gene))==1:
            zongshu='本次受检样本中共检出 %s 个纯合变异，均存在于%s基因上，'%(str(len(hom_gene)),hom_gene_set[0])
        else:
            gene='、'.join(hom_gene)
            zongshu='本次受检样本中共检出 %s 个纯合变异，分别存在于%s基因上，'%(str(len(hom_gene)),gene)
    elif len(hom_gene)==0:  ##全为杂合突变
        if len(het_gene)==1:
            zongshu='本次受检样本中共检出 1 个杂合变异，存在于%s基因上，'%(het_gene[0])
        elif len(set(het_gene))==1:
            zongshu='本次受检样本中共检出 %s 个杂合变异，均存在于%s基因上，'%(str(len(het_gene)),het_gene_set[0])
        else:
            gene='、'.join(het_gene)
            zongshu='本次受检样本中共检出 %s 个杂合变异，分别存在于%s基因上，'%(str(len(het_gene)),gene)
    else:
        zongshu='本次受检样本中共检出 %s 个变异，%s个杂合变异和%s个纯合变异，分别存在于%s和%s基因上，'%(str(len(het_gene)+len(hom_gene)),str(len(het_gene)),str(len(hom_gene)),'、'.join(het_gene),'、'.join(hom_gene))
    zongshu += get_level(hom_level,het_level)
    zongshu += '与上述基因相关的疾病及检出位点详细信息请见下表:'
    cover_Dic['detect_summary'] = zongshu
    return cover_Dic

def tongji_hethom(site_Dic):
    hom_gene = []
    het_gene = []
    hom_level = []
    het_level = []
    
    for key in sorted(site_Dic,key=lambda x:site_Dic[x][0][10]):
        value=site_Dic[key]
        siteDIc=[]
        for line in value:
            #print ('#'+line[14])
            if line[14] not in siteDIc:
                if line[4] == "纯合":
                    hom_gene.append(key)
                    hom_level.append(line[6])
                else:
                    het_gene.append(key)
                    het_level.append(line[6])
            siteDIc.append(line[14])
    return hom_gene,hom_level,het_gene,het_level

def get_level(hom,het):
    zong=hom+het
    Filt=set(zong)
    filt=[i for i in Filt]
    miaoshu=''
    if len(zong)==1:
        miaoshu = '严格遵循 ACMG 解读标准，判定为%s。'%zong[0]
    elif len(filt)==1:
        miaoshu = '严格遵循 ACMG 解读标准，均判定为%s。'%filt[0]
    elif len(het)==0:
        hom_miaoshu='、'.join(hom)
        miaoshu = '严格遵循 ACMG 解读标准，依次判定为%s。'%hom_miaoshu
    elif len(hom)==0:
        het_miaoshu='、'.join(het)
        miaoshu = '严格遵循 ACMG 解读标准，依次判定为%s。'%het_miaoshu
    else:
        hom_miaoshu='、'.join(hom)
        het_miaoshu='、'.join(het)
        miaoshu = '严格遵循 ACMG 解读标准，依次判定为%s和%s。'%(het_miaoshu,hom_miaoshu)
    return miaoshu


def get_disease_table(Dic,siteDic,BZDic):
    #疾病名称、遗传模式、疾病描述、致病性等级对应的数字
    cover_Dic={}
    tabledis=[]
    mode = prepare_zhushi(Dic)
    Dic_new=Dic
    sorted_Dic=siteDic
    miaoshubz=''
    flag=0;
    for key in sorted(sorted_Dic,key=lambda x:Dic_new[x][0][3]):  ##根据致病性排序
        gene_line=[]
        subdic=[]
        #print (key)
        #print (Dic_new[key][0][3])
        for line in Dic_new[key]:
            subkey=re.sub(' ','',line[0])+":"+ re.sub(' ','',line[1]) #该基因出现多次的疾病不再输出
            if subkey not in subdic:
                lin={'n1':line[0],'n2':line[1]}
                gene_line.append(lin)
                subdic.append(subkey)
        keylin={'n0':key,'area':gene_line}
        tabledis.append(keylin)
        if key in BZDic:
            #if flag==0:
                  #miaoshubz='注：'
            flag=1 
            joina='、'.join(BZDic[key])
            miaoshubz+='%s基因可能还与%s相关，'%(key,joina)
    if flag==1:
         miaoshubz+='但目前研究较少，如已经出现该病的临床表型，请结合临床考虑该位点的致病性。'
    cover_Dic['tbl_contents']=tabledis
    cover_Dic['result_zhujie']=miaoshubz
    #print (mode)
    cover_Dic['genetic_model']=get_model(mode)
    print (cover_Dic['genetic_model'])
    return cover_Dic



def prepare_zhushi(Dic):
    # [疾病亚型，遗传模式,疾病描述，致病等级
    
    genetic_mode=[]
    for key,value in Dic.items():
        for line in value:
            genetic_mode.append(line[1])   ##获取遗传模式
            
    mode_New = set(genetic_mode)
    mode_new = [i for i in mode_New]
    return mode_new

def get_model(li):
    genetic_model=''
    mode_Dic={'AD':'AD：常染色体显性遗传；',
    'AR':'AR：常染色体隐性遗传；',
    'XLD':'XLD：X 染色体连锁显性遗传；',
    'XLR':'XLR：X 染色体连锁隐性遗传；',
    'XL':'XL：X 染色体连锁遗传；',
    'AD/AR':'AD/AR：常染色体遗传；',
    'AR/AD':'AR/AD：常染色体遗传；'}
    for mode in li:
        if mode in mode_Dic:
            genetic_model += mode_Dic[mode]
        elif mode == r'不明':
            pass
        else:
            sys.stderr.write(u'遗传模式为%s，不在AD AR XLD XLR XL AD/AR AR/AD内'%mode)
            sys.exit(1)
    return genetic_model




def get_site_table(Dic,references,genoDic):
    table=[]
    pep=''
    cover_Dic={}
    sortDic=Dic
    rt=''
    #nue,acid,chr_site,exon,het/hom/,mute_type,zhibingxing,literature,gene_des,site_des
    for key in sorted(sortDic,key=lambda x:sortDic[x][0][10]):

        newsite={}
        for line in sortDic[key]:
            if line[14] not in newsite.keys():
                newsite[line[14]]=[line]
            else:
                newsite[line[14]].append(line)
        for site in sorted(newsite,key=lambda x:newsite[x][0][10]):
            line=newsite[site][0]
            gene_line=[]
            aa=line[2].strip().split(' ')
            bb=re.sub('Chr','',aa[0])
            bb=re.sub(':','_',bb)
            print (len(genoDic[bb]))
            if len(genoDic[bb])==0:
                origin='无'
            else:
                rt = RichText('')
                count=0
                for ori in genoDic[bb]:
                    if count > 0:
                        rt.add('\n')
                    count=count+1
                    ori=re.sub('^ ','',ori)
                    rt.add(ori,style='self-fontstyle-1')
                origin=rt
            print (origin)
            
            sitechr=line[2]+' '+ line[0]+' '+line[1]
            for line2 in newsite[site]:
                aa=line2[6]
                diseasename=line2[11]+'('+line2[12]+')'
                keyss=site+':'+aa
                refs=''
                if keyss in references.keys():
                    refs=references[keyss]
                if aa == '致病':
                    aa=RichText('致病',bold=True)
                lin={"n6":aa,"n9":refs,"n7":diseasename}
                gene_line.append(lin)
            lis={'n0':key,'n1':sitechr,'n3':line[3],'n4':line[4],'n5':line[15],'n8':origin,'area':gene_line}
            table.append(lis)

    cover_Dic['tbl_contents2']=table
    return cover_Dic

def get_jiexi_table(disease,site):
    site_sortDic=site
    disease_Dic=disease
    cover_Dic={}
    count=0
    sitejiexi=''
    genejiexi=''
    diseasejiexi=''
    sitejiexiarry={}
    genejiexiarry={}
    diseasenamearry=[]
    diseasejiexiarry=[]
    for key in sorted(site_sortDic,key=lambda x:site_sortDic[x][0][10]):

        for line in disease_Dic[key]:
            if line[0] not in diseasenamearry:
                diseasenamearry.append(line[0])
                dess=line[0]+':'+line[2]
                diseasejiexiarry.append(dess)
            else:
                pass
 
        newsite={}
        for line in site_sortDic[key]:
            if key not in genejiexiarry.keys():
                genejiexiarry[key]=line[8]
            if line[14] not in newsite.keys():
                newsite[line[14]]=[line]
            else:
                newsite[line[14]].append(line)

            for subsite in sorted(newsite,key=lambda x:newsite[x][0][10]):
                line=newsite[subsite][0]
                diseasesite=defaultdict(list)
                ss=key+':'+line[0]+','+line[1]
                if ss not in sitejiexiarry.keys():
                    sitejiexiarry[ss]={}
                for line2 in newsite[subsite]:
                    skey=line2[11]+' '+line2[12]
                    diseasesite[line2[9]].append(skey)
                for site_des in  diseasesite:
                    des_merge=','.join(diseasesite[site_des])
                    sitejiexiarry[ss][des_merge]=site_des
               
        #disease_li,disease_des,site_des=prepare_jiexi(site_sortDic[key],disease_Dic[key])
   
    diseasejiexi=[]
    genejiexi=[]
    sitejiexi=[]
   
    for li in diseasejiexiarry:
        lin={"des":li}
        diseasejiexi.append(lin)
    for li in genejiexiarry:
        ss=genejiexiarry[li]
        lin={"des":ss}
        genejiexi.append(lin)
    for subsite in sitejiexiarry:
        for dise in sitejiexiarry[subsite]:
            ss=dise+':'+sitejiexiarry[subsite][dise]
            lin={'dis':subsite,'des':ss}
            sitejiexi.append(lin)

    cover_Dic['jiexi_disease']=diseasejiexi
    cover_Dic['jiexi_site']=sitejiexi
    cover_Dic['jiexi_gene']=genejiexi
    return cover_Dic

def get_fubiao_table(Dic,siteDic,references):
    cover_Dic={}
    
    sortDic=siteDic
   
    sd = tpl.new_subdoc()
    numAll=0
    fuflag=0 
    for key in sorted(sortDic,key=lambda x:sortDic[x][0][6]):
        numAll=numAll+len(sortDic[key])
    if numAll > 0:
        fuflag=1
    print (numAll)  
    numAll=numAll+1
    table2 = sd.add_table(numAll, 7, style = sd.styles['myfubiao-2'])
    table2.autofit = False
    table2.style.font.size=Pt(7.5) 
    table2.alignment=WD_TABLE_ALIGNMENT.CENTER
    row_cells = table2.rows[0].cells
    widthtable=[1.6,2.8,2.5,1.02,2.37,3,3.51]
    titletable=['基因','变异位点','转录本号','合子状态','自然人群携带率gnomAD-东亚','ACMG变异评级','相关疾病（遗传模式）']
    for j in range(0,7):
        p = row_cells[j].paragraphs[0]
        run=p.add_run (titletable[j])
        run.bold = True #加粗
        run.font.size=Pt(7.5)
        if titletable[j]==u'转录本号':
            p = row_cells[j].add_paragraph()
            run=p.add_run (u'外显子编号')
            run.bold = True #加粗
            run.font.size=Pt(7.5)
        row_cells[j].width=Cm(widthtable[j])
        row_cells[j].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        row_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    
    LastLineNum=1
    sitenumber=0
    for key in sorted(sortDic,key=lambda x:sortDic[x][0][6]):
        newsite=defaultdict(list)
        newsite2=defaultdict(list)
        linNum=defaultdict(list)
        ACMGitems=[]
        ACMGitems_uniq=[]
        for line in sortDic[key]:
            ACMGitems.append(line[4])
            
            ACMGitems_uniq=get_acmg_item(ACMGitems)
        for line in sortDic[key]:
            acmgLev=checkACMG(line[4],ACMGitems_uniq)
            if line[9] not in newsite.keys():
                newsite[line[9]]=defaultdict(list)
            newsite[line[9]][acmgLev].append(line)
            newsite2[line[9]].append(line)
            if line[9] not in linNum:
                linNum[line[9]]=0
            linNum[line[9]]=linNum[line[9]]+1
        for site in sorted(newsite2,key=lambda x:newsite2[x][0][6]):
            arry=[]
            num=0 
            sitenumber=sitenumber+1
            for acmgIterm in sorted(newsite[site],key=lambda x:newsite[site][x][0][6]):
                num=num+len(newsite[site][acmgIterm])
                #arry.append(acmgIterm)
            arry=list(newsite[site].keys())
            num_arry=len(arry)
            line=newsite[site][arry[0]][0]  
       
            numS=linNum[site]+LastLineNum-1

            

            for j in range(5):
                table2.cell(LastLineNum,j).merge(table2.cell(numS,j))
                if (sitenumber % 2) == 1:
                    #print (sitenumber)
                    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="E1E6F5"/>'.format(nsdecls('w')))
                    table2.cell(LastLineNum,j)._tc.get_or_add_tcPr().append(shading_elm_1)
         
            shading_elm_1 = parse_xml(r'<w:shd {} w:fill="E1E6F5"/>'.format(nsdecls('w'))) 

            row_cells = table2.rows[LastLineNum].cells
            p = row_cells[0].paragraphs[0]
            run=p.add_run (key)
            run.font.size=Pt(7.5)
            row_cells[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for j in range(0,4):
                p = row_cells[j+1].paragraphs[0]
                run=p.add_run (line[j])
                run.font.size=Pt(7.5)
                row_cells[j+1].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                row_cells[j+1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
 
            
            LastLineNumACMG=LastLineNum
            for acmgIterm in sorted(newsite[site],key=lambda x:newsite[site][x][0][8]):
                numacmg=len(newsite[site][acmgIterm])
                line2=newsite[site][acmgIterm][0]
                row_cells = table2.rows[LastLineNumACMG].cells

                numA=LastLineNumACMG+numacmg-1
                table2.cell(LastLineNumACMG,5).merge(table2.cell(numA,5))
                if (sitenumber % 2) == 1:
                    shading_elm_1 = parse_xml(r'<w:shd {} w:fill="E1E6F5"/>'.format(nsdecls('w')))
                    table2.cell(LastLineNumACMG,5)._tc.get_or_add_tcPr().append(shading_elm_1)
                p = row_cells[5].paragraphs[0]
                run=p.add_run (line2[5])
                run.font.size=Pt(7.5)
                keyss=site+':'+line2[5]
                refs=''
                if keyss in references.keys():
                    refs=references[keyss]
                if refs != '':
                    run = p.add_run(refs)
                    run.font.size=Pt(8)
                    run.font.superscript=True
                row_cells[5].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                row_cells[5].vertical_alignment = WD_ALIGN_VERTICAL.CENTER 
                
                disline=LastLineNumACMG
                for line2 in newsite[site][acmgIterm]:
                    con=line2[7]+'('+line2[8]+')'
                    row_cells = table2.rows[disline].cells
                    if (sitenumber % 2) == 1:
                        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="E1E6F5"/>'.format(nsdecls('w')))
                        row_cells[6]._tc.get_or_add_tcPr().append(shading_elm_1)
                    p = row_cells[6].paragraphs[0]
                    run=p.add_run (con)
                    run.font.size=Pt(7.5)
                    disline=disline+1
                LastLineNumACMG=LastLineNumACMG+numacmg
            LastLineNum=LastLineNum+linNum[site]
    cover_Dic['fubiaotable']=sd
    cover_Dic['tbl_contents3']=fuflag
    
    return cover_Dic


def get_acmg_item(ACMGitem):
    Dic=[]
    for item in ACMGitem:
        l=[x.strip() for x in item.strip().split(' ')]
        l=sorted(l)
        sortl=' '.join(l)
        if sortl not in Dic:
            Dic.append(sortl)
    return Dic


def checkACMG(items,Dic):
    l=[x.strip() for x in items.strip().split(' ')]
    l=sorted(l)
    sortl=' '.join(l)
    fitem='-'
    for item in Dic:
        if item == sortl:
            fitem= item
    if fitem=='-':
        sys.stderr.write(u'请检查附表的致病证据等级')
         
    return fitem

###############################################参考文献
#referenceAll,referenceSite=getReference(module,litetable)
#[line_list[78],line_list[83],line_list[84],line_list[85],line_list[82]] 致病性，疾病名称，遗传模式，基因,参考文献
def getReference(module,litetable):

    liteAll=[]
    liteAll2={}
    litesite={}
    CS_lite=defaultdict(list)
    newlite=defaultdict(list)
    if module=='CS':
        CS_lite=read_CSdata()
    
    numberite=0;
    for eachsite in litetable:
        for line_list in litetable[eachsite]:
            key1=re.sub(' ','',line_list[3])+':'+re.sub(' ','',line_list[1])+':'+re.sub(' ','',line_list[2]) 
            key2=eachsite+':'+line_list[0]
            newlite[key2].append(line_list[4])
            if key1 in CS_lite.keys():
                line=CS_lite[key1]
                newlite[key2].append(line[0])
                newlite[key2].append(line[1])
    
    for eachsite in newlite:
        litearry=get_wenxian(newlite[eachsite])
        arry=[]
        for li in litearry:
            num=0
            if li not in liteAll2.keys():
                numberite=numberite+1
                num=numberite
                liteAll2[li]=numberite
                liteAll.append([numberite,li])
            else:
                num=liteAll2[li]
            if num !=0:
                arry.append(num)
        arry2=set(arry)
        arry2=sorted(arry2)
        laststart=0
        lastend=0
        flag=0
        arry3=[]
        for i in arry2:
            if flag==0:
                laststart=i
                lastend=i
                flag=1
            else:
                if i == lastend+1:
                    lastend=i
                elif i > lastend+1:
                    if laststart == lastend:
                        arry3.append(str(laststart))
                    else:
                        ss=str(laststart)+'-'+str(lastend)
                        arry3.append(ss)
                    laststart=i
                    lastend=i
                else:
                    sys.stderr.write('文献排序错误')
                    sys.exit(1)
        if laststart == lastend and laststart>0:
            arry3.append(str(laststart))
        elif laststart>0 :
            ss=str(laststart)+'-'+str(lastend)
            arry3.append(ss)
        sssss=','.join(arry3)
        if sssss != '':
            litesite[eachsite]='['+sssss+']'
    numberite=numberite+1
    liteAll.append([numberite,'Richards Sue,Aziz Nazneen,Bale Sherri,Bick David,Das Soma,Gastier-Foster Julie,Grody Wayne W,Hegde Madhuri,Lyon Elaine,Spector Elaine,Voelkerding Karl,Rehm Heidi L,ACMG Laboratory Quality Assurance Committee.Standards and guidelines for the interpretation of sequence variants: a joint consensus recommendation of the American College of Medical Genetics and Genomics and the Association for Molecular Pathology.[J].Genet. Med.2015 May;17(5):405-24.'])
    return liteAll,litesite


def read_CSdata():
    data='/PUBLIC/pipline/script/Report/XKFW/V3_180824/database/cs.data4.txt'
    disease_Dic=defaultdict(list)
    if os.path.isfile(data) == False:
        sys.stderr.write('%s is not a file' % data)
        sys.exit(1)
    try:
        with open(data, 'r+',encoding='utf-8',errors='ignore') as file_handle:
            for line in file_handle:
                if re.match('^疾病大类',line.strip()):continue
                line_1=re.sub('%','%',line)
                #line_2=re.sub(r'\',r'\\',line_1)
                line_2=re.sub('_','_',line_1)
                line_list=[x.strip() for x in line_2.strip().split('\t')]
                if len(line_list) < 9:
                    sys.stderr.write(u'cs.data.txt 列数不足 ')
                    sys.stderr.write(line_2)
                    sys.exit(1)
                disease_key=re.sub(' ','',line_list[1])+':'+re.sub(' ','',line_list[3])+':'+re.sub(' ','',line_list[4])     #gene:disease_subtype:遗传模式 作为key
                disease_Dic[disease_key]=[line_list[7],line_list[8]]  # A类文献，B类文献
    except Exception as e:
        traceback.print_exc(e)
    return disease_Dic

def get_wenxian(li):
    wenxian=[]
    wenxian_new=[]
    for i in li:
        if i=='' or i =='-':
            pass
        else:
            wenxian.append(i)
    wenxian_new=prepare_lite(wenxian)
    return wenxian_new


def prepare_lite(li):
    li_new=[]
    for i in li:
        if '|' in i:
            arr=i.split('|')
            for jj in arr:
                if jj=='' or jj =='-' or jj =='.':
                    pass
                else:
                    li_new.append(jj)
                #li_new += arr
        else:
            li_new.append(i)
    return li_new

############################################################3

def get_QC_Inf(QCfile,sampleid):
    QC_inf_list=[]
    QC_li=[]
    cover_Dic= defaultdict(list)
    head=['Q30','cover_target','cover_20','depave']
    if os.path.isfile(QCfile) == False:
        print('%s is not a file' % QCfile)
        return cover_Dic
    try:
        file_handle = open(QCfile, 'rU')
        QC_list = file_handle.readlines()
        samq30=''
        for line in QC_list:
            QC_arr = [x.strip() for x in line.strip().split('\t')]
            if QC_arr[0] == sampleid and len(QC_arr) > 40:
                QC_li.append(QC_arr)
            elif QC_arr[0] == sampleid and len(QC_arr) < 40:
                samq30=QC_arr[10]+';'+QC_arr[11]
            else:
                pass
        if len(QC_li)==1:
            QC_inf = [samq30,QC_li[0][2],QC_li[0][32],QC_li[0][3]]
        else:
            sys.stderr.write(u'请检查质控结果是否无或未合并')
            sys.exit(1)
        cover_Dic = {k:v for k,v in zip(head,QC_inf)}
    except Exception as e:
        traceback.print_exc(e)
    finally:
        if 'fileHandle' in dir() or file_handle.closed == False:
            file_handle.close()
    return cover_Dic


def get_HTB(sampleid):
    #files=glob.glob(r'/PUBLIC/pipline/database/sheet_hutong/*心康互通表*.xls*')
    files=glob.glob(r'/PUBLIC/pipline/database/sheet_hutong/2019心康送样信息单*.xls*')
    max_num=0
    age=date=''
    sex=''
    age_num = '' #wph changed 181106
    date_num = ''
    shizinum = ''
    for i in files:
        hutong=i.split('/')[-1]
        #st=re.match(r'心康互通表（([0-9]*)）.xlsx',hutong)
        st=re.match(r'2019心康送样信息单-([0-9]*).xlsx',hutong)
        if int(st.group(1))> max_num:
            max_num=int(st.group(1))
    HTB=r'/PUBLIC/pipline/database/sheet_hutong/2019心康送样信息单-%s.xlsx' %max_num
    try:
        data = xlrd.open_workbook(HTB)
        table = data.sheet_by_name('下单表')
        nrows = table.nrows
        for rownum in range(1,nrows):
            if table.row(rownum)[2].value == sampleid:
                age=table.row(rownum)[8].value
                sex=table.row(rownum)[7].value
                doc=table.row(rownum)[10].value
                stype=table.row(rownum)[11].value
                shizinum=table.row(rownum)[3].value
                print (shizinum)
                ctype = table.cell(rownum, 1).ctype
                cell =table.cell_value(rownum, 1)
                date
                if ctype ==3 :
                    date = datetime(*xldate_as_tuple(cell, 0))
                    date = date.strftime('%Y/%m/%d')
                else:
                    date=table.row(rownum)[1].value
        date_num=date.replace("/","-")
        if isinstance(age, float) == True: #wph changed 181106
            age_num=str(int(age)) + "岁"
        elif not age:
            age_num="不详"
        else:
            age_num=age
        if age_num=='' or date_num=='':
            sys.stderr.write(u"ID不在互通表 或 互通表中没有年龄/送样日期信息")
            exit
    except Exception as e:
        traceback.print_exc(e)
    return age_num,date_num,sex,doc,stype,shizinum

def get_sampleInf(sampleid,cfgFile,feature_file):
    sample_inf_list=[]
    cover_Dic = defaultdict(list)
    head=['sampleID','name','sex','age','data1','data2','chenghu','clinical','doctor','sampletype','jiazushi','sampleIDNF']
    time2str=time.strftime('%Y-%m-%d',time.localtime(time.time()))
    if os.path.isfile(cfgFile) == False:
        print('%s is not a file' % cfgFile)
        return cover_Dic
    try:
        with open(cfgFile, 'rU') as file_handle:
            for line in file_handle:
                line_list=[x.strip() for x in line.strip().split('\t')]
                if line_list[0]==sampleid:
                    line_list[5]=re.sub(r'\(.*?\)','',line_list[5])
                    (age,date,sex,doc,samtyp,shizinum)=get_HTB(sampleid) #wph changed 190423 修改性别不匹配的问题
                    if sex not in line_list[5]:
                        line_list[5] = sex
                    chenghu=judge_chenghu(line_list[5])
                    feature,familyhistory=get_feature(feature_file,sampleid)
                    sample_inf_list=[line_list[0],line_list[4],line_list[5],age,date,time2str,chenghu,feature,doc,samtyp,familyhistory,shizinum]  #需要送样日期，暂时用出报告日期代替
                    break
        cover_Dic = {k:v for k,v in zip(head,sample_inf_list)}
    except Exception as e:
        traceback.print_exc(e)
    return cover_Dic


def judge_chenghu(sex):
    chenghu=''
    try:
        if sex ==r'女':
            chenghu=r'女士'
        elif sex ==r'男':
            chenghu=r'先生'
        else:
            print('性别是%s,check please' % sex)
    except Exception as e:
        traceback.print_exc(e)
    return chenghu


def get_feature(feature_file,sampleid):
    familyhistory=''
    feature = '无临床信息'
    if os.path.isfile(feature_file) == False:
        print('%s is not a file' % feature_file)
        return 'Error File'
    try:
        file_handle = open(feature_file, 'rU')
        feature_list = file_handle.readlines()
        for line in feature_list:
            feature_arr = [x.strip() for x in line.strip().split('\t')]
            if feature_arr[0] == sampleid and len(feature_arr)>1:
  
                pattern = re.compile(r'(临床诊断疾病: [^\\]+)', re.I)
                m = re.findall(pattern,feature_arr[1])
                if m is None or len(m)==0:
                    feature2='-'
                else:
                    feature2=m[0]
                    feature2=feature2.strip().strip(',|，')
                    feature=feature2
            
                pattern = re.compile(r'(临床症状及发病年龄: [^\\]+)', re.I)
                #m = pattern.match(feature_arr[1])
                m = re.findall(pattern,feature_arr[1])
                if m is None or "临床症状及发病年龄" not in feature_arr[1]:
                    feature1='-'
                else:
                    #feature=r''' %s ''' %(m.groups(0)[0])
                    feature1=m[0]
                    feature1=feature1.strip().strip(',|，')
                    if feature == '无临床信息':
                        feature=feature1
                    else:
                        feature=feature+'; '+feature1

                pattern = re.compile(r'何人（与患者关系）: ([^\\]+)', re.I)
                #m = pattern.match(feature_arr[1])
                m = re.findall(pattern,feature_arr[1])
                
                if m is None or len(m)==0 or "与患者关系" not in feature_arr[1]:
                    familyhistory='-'
                else:
                    #familyhistory=r''' %s ''' %(m.groups(0)[0])
                    print (m[0])
                    familyhistory=m[0]
                    familyhistory=familyhistory.strip().strip(',|，')
            else:
                pass
    except Exception as e:
        traceback.print_exc(e)
    print (feature)
    print (familyhistory)
    return feature,familyhistory


def get_CS():
    cover_Dic={}
    CS_li=[]
    head_Dic=[]
    tabledis=[]
    gene_line=['xinlv','xinji','zhudongmai','xinzang','qita']
    head_Dic={'xinlv':'遗传性心律失常（59个）',
    'xinji':'遗传性心肌病（101个）',
    'zhudongmai':'遗传性主动脉病（12个）',
    'xinzang':'先天性心脏病（24个）',
    'qita':'其他心脏相关疾病（12个）'
    }
    for disbig in gene_line:
        disname=head_Dic[disbig]
        CS_file='/PUBLIC/pipline/script/Report/XKFW/V3_180824/database/cs_disease_'+ disbig+'.txt'
        Dic_new=read_CS(CS_file)
        gene_line=[]
        for line in Dic_new:
            lin={'n2':line[0],'n3':line[2]}
            gene_line.append(lin)
        keylin={'n1':disname,'area':gene_line}
        tabledis.append(keylin)

    cover_Dic['tbl_gene']=tabledis
    return cover_Dic

def get_GXY():
    cover_Dic={}
    
    tabledis=[]
    filen='/PUBLIC/pipline/script/Report/XKFW/V3_180824/database/gxy_disease_name3.txt'
    Dic_new=read_CS(filen)
    gene_line=[]
    for line in Dic_new:
        lin={'n1':line[0],'n2':line[2]}
        tabledis.append(lin)

    cover_Dic['tbl_gene_gxy']=tabledis
    return cover_Dic


def read_CS(CS_file):
    CSxiang=[]
    if os.path.isfile(CS_file) == False:
        print('%s is not a file' % CS_file)
        return CSxiang
    try:
        with open(CS_file, 'rU') as file_handle:
            for line in file_handle:
                line_list=[x.strip() for x in line.strip().split('\t')]
                line_list_col=[x.strip() for x in line_list[0].strip().split('\\\\')]
                li=[line_list_col[0],line_list_col[1],line_list[1]]
                CSxiang.append(li)

    except Exception as e:
        traceback.print_exc(e)
    return CSxiang

def get_IDT(filen):
    cover_Dic={}
    tabledis=[]
    Dic_new=[]
    gene_number=0
    disease_name=''
    gene_dic=[]
    if os.path.isfile(filen) == False:
        print('%s is not a file' % filen)
        return Dic_new
    try:
        with open(filen, 'rU') as file_handle:
            i=0
            for line in file_handle:
                line_list=[x.strip() for x in line.strip().split('\t')]
                if line_list[0]==sampleid:
                    gdic=line_list[2].strip().split(', ')
                    for gene in gdic:
                        gene_dic.append(gene)
                    if i==0 and  line_list[1] !=u'高血压':
                        disease_name=line_list[1]
                    elif i>0 and  line_list[1] !=u'高血压':
                        disease_name=disease_name+','+line_list[1]
                    if line_list[1]==u'心源性猝死':
                        line_list[2]='ABCA3, ABCC9, ACTA1, ACTA2, ACTC1, ACTN2, ACVRL1, AKAP9, ALDH1A2, ALG10, ANK2, ANKRD1, APOB, BAG3, BMPR2, BRAF, BVES, CACNA1C, CACNA2D1, CACNA2D4, CACNB2, CALM1, CALM2, CALM3, CALR3, CASQ2, CAV3, CHRM2, COL3A1, COL6A1, COL6A2, COL6A3, CRELD1, CRELD2, CRYAB, CSRP3, CTF1, CTNNA3, DES, DMD, DOLK, DPP6, DSC2, DSG2, DSP, DTNA, EIF2AK4...'
                    li=[line_list[1],line_list[2]]
                    Dic_new.append(li)
                    i=i+1      
    
    except Exception as e:
        traceback.print_exc(e)
    gene_line=[]
    for line in Dic_new:
        lin={'n1':line[0],'n2':line[1]}
        tabledis.append(lin)

    gene_number=len (set(gene_dic))    
    disease_name=re.sub('单基因糖尿病,','',disease_name)
    cover_Dic['tbl_gene_IDT']=tabledis
    cover_Dic['diesanse_table_name']=disease_name
    cover_Dic['gene_table_number']=gene_number
    return cover_Dic

def read_sample_info_PCR(figfile):
    het_hom = {'00':'野生纯合','01':'变异杂合','11':'变异纯合'}
    sex=read_sample_sex()
    samples=[]
    sitetable=[]
    sangerperson={}
    site_genotype={}
    site_genotype2=defaultdict(list)
    sitetables={}
    genotype_hash={}

    try:
        with open(figfile, 'rU') as file_handle:
            for line in file_handle:
                line_list=[x.strip() for x in line.strip().split('\t')]
                if line_list[15] == 'NULL':
                    continue
                if line_list[1] == sampleid and line_list[2] not in samples:
                        line_list[0]=re.sub('妈妈','母亲',line_list[0])
                        line_list[0]=re.sub('爸爸','父亲',line_list[0])
                        ssex=sex[line_list[2]]
                        if line_list[0] not in sangerperson.keys():
                            sangerperson[line_list[0]]=[[line_list[2],line_list[11],line_list[10],line_list[0],line_list[12],ssex]] 
                                                         # 2 心康编号 11 拭子编号 10 受检者姓名 0 家系关系 12 样本类型 性别
                        else:
                            sangerperson[line_list[0]].append([line_list[2],line_list[11],line_list[10],line_list[0],line_list[12],ssex])
                        samples.append(line_list[2])
                sitesub=line_list[3]+'_'+line_list[4]
                if line_list[2] == sampleid and  line_list[2]!='-' and  sitesub not in sitetables.keys():
                        site_type=het_hom[line_list[15]]
                        ssex=sex[line_list[2]]
                        if line_list[3] and "X" and ssex == "男" and site_type == "变异纯合":
                            site_type = "半合子";
                        if site_type not in genotype_hash.keys():
                            genotype_hash[site_type]=0
                        genotype_hash[site_type]=genotype_hash[site_type]+1
                        sitetables[sitesub]=line_list[6] #赋值为基因名字
                        sitetable.append([line_list[3],line_list[4],line_list[5],line_list[6],line_list[7],line_list[8]]) ##chr,start,rsnum,gene,cds,pep

                if line_list[1] == sampleid:
                    site_type=het_hom[line_list[15]]
                    if line_list[3] and "X" and ssex == "男" and site_type == "变异纯合":
                        site_type = "半合子";
                    if line_list[2] not in site_genotype.keys():
                        site_genotype[line_list[2]]={}
                    site_genotype[line_list[2]][sitesub]=site_type
                    if line_list[2] != sampleid:
                        sssss=line_list[0]+site_type
                        site_genotype2[sitesub].append(sssss)

    except Exception as e:
        traceback.print_exc(e)

    return sangerperson,sitetable,site_genotype,genotype_hash,sitetables,site_genotype2


def read_sample_sex():
    files=glob.glob(r'/PUBLIC/pipline/database/sheet_hutong/2019心康送样信息单*.xls*')
    max_num=0
    date_num = ''
    for i in files:
        hutong=i.split('/')[-1]
        st=re.match(r'2019心康送样信息单-([0-9]*).xlsx',hutong)
        if int(st.group(1))> max_num:
            max_num=int(st.group(1))
    HTB=r'/PUBLIC/pipline/database/sheet_hutong/2019心康送样信息单-%s.xlsx' %max_num
    print (HTB)
    id=sampleid;
    id=re.sub('^N(0+)','',id)
    sample_sex={}
    try:
        data = xlrd.open_workbook(HTB)
        table = data.sheet_by_name('下单表')
        nrows = table.nrows
        guanxis=[]
        for rownum in range(1,nrows):
            tag=str(table.row(rownum)[6].value)
            tag2=str(table.row(rownum)[6].value)
            if table.row(rownum)[2].value == sampleid:
                sample_sex[sampleid]=table.row(rownum)[7].value
            elif u'一代验证' in tag :
                pattern = re.compile(r'一代验证-(\d+)-(.*)', re.I)
                m = pattern.match(tag2)
                number=''
                if m is None:
                    number='nnn'
                else:
                    number=r''' %s ''' %(m.groups(0)[0])
                    number=re.sub(' ','',number)
                if number==id:
                    if m.groups(0)[1] is None:
                        sys.stderr.write('%s is not a right sanger type' % table.row(rownum)[6])
                        sys.exit(1)
                    guanxi=r''' %s ''' %(m.groups(0)[1])
                    guanxi=re.sub('妈妈','母亲',guanxi)
                    guanxi=re.sub('爸爸','父亲',guanxi)
                    sample_sex[table.row(rownum)[2].value]=table.row(rownum)[7].value
    except Exception as e:
        traceback.print_exc(e)
    return sample_sex


def get_sanger_Inf_Family(gene_strand):
    genotype_hash=['野生纯合','变异杂合','变异纯合','半合子'];
    path1=os.path.abspath(filterFile)
    path_mod ={'GXY':u'01GXY','CS':u'02Cus','IDT':u'06IDT'}
    if module == 'GXY':
        print ('ddd')
        path1=re.sub('01products/'+path_mod[module]+'/'+'GXY_sample_rigth_to_PDF_final.txt','',path1)
    else:
        path1=re.sub('01products/'+path_mod[module]+'/'+module+'.pdf.txt','',path1)
    path1=path1+'02individual/'+sampleid+'/fig/'
    fig=path1+sampleid+'.family.site.txt';
    cover_Dic={}
    info_line=[]
    info_line2=[]
    sitDic=[]
    # col_label=[]
    conclude={}
    conclude['descript']=''
    site_genotype2=defaultdict(list)
    family_sign=0
    print (fig)
    if os.path.exists(fig):   

        (sangerpersoninfo,siteDic,siteGenotype,genotype_hash,sitetables,site_genotype2)=read_sample_info_PCR(fig)
        all_sample=[]
        cols=[]
        relat={}
        for key in sangerpersoninfo:
            if key=='先证者':
                for line in sangerpersoninfo[key]:
                    #line[2]=''
                    all_sample.append(line[0])
                    relat[line[0]]=[line[3],line[2]]
        for key in sangerpersoninfo:
            if key != '先证者':
                family_sign=1
                for line in sangerpersoninfo[key]:
                    all_sample.append(line[0])
                    relat[line[0]]=[line[3],line[2]]

        conclude=sangertexconclu(sangerpersoninfo,siteDic,siteGenotype,genotype_hash,sitetables,family_sign)
        
                        
        for line in siteDic:
            samplesigh=[]
            count=0
            site_sample=line[0]+'_'+line[1]
            site_sample2=line[0]+'-'+line[1]
            strandsite=gene_strand[site_sample2]
            fangxiang=''
            if strandsite=='+':
                fangxiang='(正向)'
            elif strandsite=='-':
                fangxiang='(负向)'
            siteu=u'检测位点：'+line[3]+':'+line[4]+','+line[5]
            print (siteu)
            for sam in all_sample:
                if sam == sampleid:
                    relat[sam][0]=''
                figname=sam+'-'+line[0]+'-'+line[1]+'.jpg'
                figname=re.sub('Chr','',figname)
                figname=re.sub(':','-',figname)
                figpath=path1+figname
                genosub=siteGenotype[sam][site_sample]
                print (figpath)
                ss=u'先证者'+relat[sam][0]+'（'+relat[sam][1]+'）'
                siteu2=siteu+','+genosub+fangxiang
                lis={}
                print(ss)
                if os.path.exists(figpath):
                    lis={'sample':ss,'site':siteu2,'con':InlineImage(tpl,figpath,width=Mm(179),height=Mm(47)),'space1':'\n','space2':'\n'}
                else:
                    lis={'sample':ss,'site':siteu2,'con':'参照序列\n样本序列\n\n\n\n','space1':'\n','space2':'\n'}
                       
                sitDic.append(lis)
   
     
    #cover_Dic['col_labels']=col_label
    #cover_Dic['sangerSam']=info_line
    #cover_Dic['sangerResult']=info_line2
    cover_Dic['family_sign']=family_sign
    cover_Dic['varientFamily']=sitDic
    cover_Dic['descript']=conclude['descript']
    print (site_genotype2)
    return  cover_Dic,site_genotype2

def sangertexconclu(sangerperson,sitetable,siteGenotype,genotype_hash,sitetables,family_sign):
    conclu=''
    concluall=''
    genotyp_count=''
    geneall=''
    mutation=''
    unmutation=''
    father_flag=0
    father_ID=''
    mother_flag=0
    mother_ID=''
    count={}
    family=[];
    genotype_list=genotype_hash
    siteList=sitetables
    geneDic=[]
    persen=sangerperson
    siteDic=sitetable
    propositusID=''
    conclude={}


    for key in genotype_list:
        if genotype_list[key]:
            genotyp_count=genotyp_count + str(genotype_list[key])+u'个'+key+ '，'
    print(genotyp_count)
    print ('#')
    genotyp_count=re.sub('变异杂合','杂合变异',genotyp_count)
    genotyp_count=re.sub('变异纯合','纯合变异',genotyp_count)
    genotyp_count=re.sub('半合子','半合子变异',genotyp_count)
    
    for key in siteList:
        geneDic.append(siteList[key])

    set(geneDic)
    
    geneall='、'.join(geneDic)

    if family_sign==1:
        concluall='该样本中检测出 %s 存在于 %s 基因上。家系验证结果显示:'%(genotyp_count,geneall)
    rt = RichText(concluall)

    site_count=len(siteList)

    propositusID=persen['先证者'][0][0]
    for key in persen:
        if key not in family:
            family.append(key)
        if key == "母亲":
            mother_flag=len(persen['母亲'])
            mother_ID=persen['母亲'][0][0]
            if mother_flag>1:
                sys.stderr.write('%s 母亲数目大于1' % mother_flag)
                sys.exit(1)
                

        if key == "父亲":
            father_flag=len(persen['父亲'])
            father_ID=persen['父亲'][0][0]
            if father_flag>1:
                sys.stderr.write('%s 父亲数目大于1' % mother_flag)
                sys.exit(1)

    print ("位点数目"+str(site_count)+"\n")

    numbersite=len(siteDic)
    countsite=0
    for l in siteDic:
        #rt.add('\n')
        countsite=countsite+1
        concluall=''
        if numbersite >1:
                concluall='('+str(countsite)+')'
        concluall=concluall+" %s 基因上的变异（ %s ，%s ）" %(l[3],l[4],l[5])
        sitesub=l[0]+'_'+l[1]
        mutys=[];
        unmutys=[];

        if mother_flag==0 and father_flag==0:
            concluall=concluall+"，";
            conclu=u'由于先证者的父母未验证，无法确定变异的遗传来源。';
        elif mother_flag==1 and father_flag==1:
            print (siteGenotype[propositusID][sitesub])
            print (siteGenotype[mother_ID][sitesub])
            print (siteGenotype[father_ID][sitesub])
            if siteGenotype[propositusID][sitesub] == siteGenotype[mother_ID][sitesub] and siteGenotype[propositusID][sitesub] != siteGenotype[father_ID][sitesub] or siteGenotype[propositusID][sitesub] == "半合子" and siteGenotype[mother_ID][sitesub] == '变异杂合':
                concluall=concluall+"，先证者的母亲 %s 携带该变异，"%(persen['母亲'][0][2]);
                concluall=concluall+"父亲 %s 未携带该变异，"%(persen['父亲'][0][2]);
                conclu="可见先证者的变异遗传自母亲。";
            elif  siteGenotype[propositusID][sitesub] == siteGenotype[father_ID][sitesub]:
                concluall=concluall+"，先证者的父亲 %s 携带该变异，"%(persen['父亲'][0][2]);
                concluall=concluall+"母亲 %s 未携带该变异，"%(persen['母亲'][0][2]);
                conclu="可见先证者的变异遗传自父亲。";            
            else:
                concluall=concluall+"父亲 %s 、母亲 %s 未携带该变异，"%(persen['父亲'][0][2],persen['母亲'][0][2]);
                conclu=u'可见先证者的变异为denovo变异。';
        elif  mother_flag==1 and  father_flag==0:
            if siteGenotype[propositusID][sitesub] == siteGenotype[mother_ID][sitesub]  or siteGenotype[propositusID][sitesub] == "半合子" and siteGenotype[mother_ID][sitesub] == "变异杂合":
                concluall=concluall+"，先证者的母亲 %s 携带该变异，"%(persen['母亲'][0][2]);
            else:
               concluall=concluall+"，先证者的母亲 %s 未携带该变异，"%(persen['母亲'][0][2]);
               conclu=u'由于父亲未验证，无法确定变异的遗传来源。';
        elif  mother_flag==0 and  father_flag==1:
            if  siteGenotype[propositusID][sitesub] == siteGenotype[father_ID][sitesub]:
                concluall=concluall+"，先证者的父亲 %s 携带该变异，"%(persen['父亲'][0][2]);
            else:
                concluall=concluall+"，先证者的父亲 %s 未携带该变异，"%(persen['父亲'][0][2]);
                conclu=u'由于母亲未验证，无法确定变异的遗传来源。';



        for relation in family:

            if relation=="先证者" or relation=="母亲" or relation=="父亲":
                continue
            for li in persen[relation]:
                rid=li[0]
                print (rid)
                print (sitesub)
                if siteGenotype[rid][sitesub] == siteGenotype[propositusID][sitesub]:
                    mutys.append(relation+li[2])
                else:
                    unmutys.append(relation+li[2])

        if len(mutys) != 0:
            mutation='、'.join(mutys)+"携带该变异，"
        if len(unmutys) != 0:
            unmutation='、'.join(unmutys)+"未携带该变异，";

        concluall=concluall+mutation+unmutation+conclu
        rt.add(concluall,style='self-fontstyle-1')
        
        mutation='';
        unmutation='';
    #concluall=concluall+u'具体请结合临床分析，以下为本次检测样本的测序峰图。';
    #rt.add(u'具体请结合临床分析，以下为本次检测样本的测序峰图。',style='self-fontstyle-1')
    if family_sign==0:
        rt='该样本中检测出 %s 存在于 %s 基因上。由于未经验证，无法确定变异的遗传来源.'%(genotyp_count,geneall)
    conclude['descript'] = rt
    print (rt)
    return conclude



def main():
	argInit()
	context={}
	referencearr=[]

	#tpl=DocxTemplate('南方医院-基因检测报告模版20190517.docx')

	id_filter,id_filterfu,id_filterbz = get_filter(filterFile,sampleid)
	diseaseDic,siteDic,diseaseDic_fu,siteDic_fu,miaoshubzz,litetable,gene_strand = prepare_allDic(id_filter,id_filterfu,id_filterbz) #wph add 190325
	referenceAll,referenceSite=getReference(module,litetable)
	zongshu = get_zongshu(siteDic)
	disease_table = get_disease_table(diseaseDic,siteDic,miaoshubzz)
	SangerFamily,site_genotype2=get_sanger_Inf_Family(gene_strand)
	site_table = get_site_table(siteDic,referenceSite,site_genotype2)
	fubiao_table = get_fubiao_table(diseaseDic_fu,siteDic_fu,referenceSite)
	jiexi_table = get_jiexi_table(diseaseDic,siteDic)
	sampleInf_Dic = get_sampleInf(sampleid,cfgFile,feature_file)
	QC_Inf_Dic = get_QC_Inf(QCfile,sampleid)
	#Sangersite = get_sanger_Inf(siteDic)

	context['detect_summary']=zongshu['detect_summary']
	context['tbl_contents']=disease_table['tbl_contents']
	context['result_zhujie']=disease_table['result_zhujie']
	
	context['tbl_contents2']=site_table['tbl_contents2']
	context['tbl_contents3']=fubiao_table['tbl_contents3']
	context['fubiaotable']=fubiao_table['fubiaotable']
	context['genetic_model']=disease_table['genetic_model']
	context['vartioninfo']=jiexi_table['jiexi_site']
	context['diseaseinfo']=jiexi_table['jiexi_disease']
	context['geneinfo']=jiexi_table['jiexi_gene']
	
	context['family_sign']=SangerFamily['family_sign']
	context['varientFamily']=SangerFamily['varientFamily']
	context['descript']=SangerFamily['descript']
		
	lite=[]
	for item in referenceAll:
		indexs=item[0]
		indexsa=str(indexs)+'.'
		refline={'index':indexsa,'con':item[1]}
		lite.append(refline)
	context['reference']=lite

	for key in QC_Inf_Dic:
            context[key]=QC_Inf_Dic[key]
	
	for key in sampleInf_Dic:
            context[key]=sampleInf_Dic[key]
 	
	context['telphone']='-'
	#context['break']=tpl.add_page_break()
	context['program']='遗传性心血管基因检测Panel'
	context['testMethmod']='二代基因测序'
	context['laboratory']='-'
	context['reportHeader']=context['name']
	#context['sampleIDNF']='JK21004B191224007' 
        
	if module == 'CS':
	    context['CS_moudle']=1
	    geneListtable=get_CS()
	    context['tbl_gene']=geneListtable['tbl_gene']
	    context['program']='遗传性心血管基因检测Panel'
	elif module == 'GXY':
	    context['GXY_moudle']=1
	    geneListtable=get_GXY()
	    context['tbl_gene_gxy']=geneListtable['tbl_gene_gxy']
	    context['program']='单基因高血压/血钾异常Panel'
	elif module == 'IDT':
	    context['IDT_moudle']=1
	    geneListtable=get_IDT(GeneListfile)
	    context['tbl_gene_IDT']=geneListtable['tbl_gene_IDT']
	    context['diesanse_table_name']=geneListtable['diesanse_table_name']
	    context['gene_table_number']=geneListtable['gene_table_number']

	    context['program']='遗传病性疾病全外显子基因检测'
	programs=context['program']
	programs=re.sub('/','-',programs)
	sd2 = tpl.new_subdoc()  
	sd2.add_page_break()
	context['break']=sd2
	tpl.render(context,autoescape=True)
	tpl.save(outDir + "/" + sampleid+'.'+programs+'报告.docx')
	#set_updatefields_true(sampleid+'.遗传性心血管基因检测Panel报告.docx')

if __name__ == '__main__':
	main()




