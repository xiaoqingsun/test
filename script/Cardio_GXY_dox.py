#-*- coding:utf-8 -*-
import os
import sys
#reload(sys)
#sys.setdefaultencoding("utf-8")
import shutil
import logging
import argparse
import traceback
import time
import decimal
import lxml
import xlrd
import os
import glob
import sys
import re
from string import Template
from collections import defaultdict
from docx import Document
from docx.shared import Mm, Inches, Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Cm,Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docxtpl import DocxTemplate
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docxtpl import DocxTemplate, InlineImage, RichText

tpl=DocxTemplate('/PUBLIC/pipline/script/Report/XKFW/V3_180824/tempt/南方医院报告模板-高血压.docx')

def argInit():
    parser = argparse.ArgumentParser(
        description="report_create program.")
    parser.add_argument(
        '-site', help='file GXY.filter.txt', required=True)
    parser.add_argument(
        '-part', help='2 / 3 / 4', required=True)
    parser.add_argument(
        '-out', help='project name', default=os.getcwd())
    parser.add_argument(
        '-sampleID', help='sampleid', required=True)
    parser.add_argument(
        '-disease', help='GXY.disease.table', required=True)
    parser.add_argument(
        '-module', help='module', required=True)
    parser.add_argument(
        '-yinyang', help='阴阳性', required=True)
    parser.add_argument(
        '-cfg', help='file cfg.txt', required=True)
    parser.add_argument(
        '-feature', help='feature', required=True)
    parser.add_argument(
        '-qc', help='file Stat_QC.xls', required=True)
    argv = vars(parser.parse_args())
    global filterFile, part,outDir,sampleid,disease_file,module,yinyang_file,cfgFile,QCfile,feature_file
    filterFile = argv['site'].strip()
    part = argv['part'].strip()
    outDir = os.path.abspath(argv['out'].strip())
    sampleid=argv['sampleID'].strip()
    disease_file=argv['disease'].strip()
    module=argv['module'].strip()
    yinyang_file=argv['yinyang'].strip()
    cfgFile = argv['cfg'].strip()
    QCfile = argv['qc'].strip()
    #sampleid = argv['sample'].strip()
    feature_file = argv['feature'].strip()

def get_filter(filterFile,sampleid):
    filter_li=[]
    filter_fu=[]
    filter_bz=[]
    id_li=[]
    if os.path.isfile(filterFile) == False:
        sys.stderr.write('%s is not a file' % filterFile)
        sys.exit(1)
    try:
        with open(filterFile, 'rU') as file_handle:
            for line in file_handle:
                if re.match('^Sample',line.strip()):continue
                if re.match('^####',line.strip()):continue
                line_list=[x.strip() for x in line.strip().split('\t')]
                if re.match('^###',line.strip()):
                    newsam=re.sub('###','',line_list[0])
                    if  newsam==sampleid:
                        filter_bz.append(line.strip())
                elif re.match('^##N',line.strip()):
                    newsam=re.sub('##','',line_list[0])
                    if  newsam==sampleid:
                        filter_fu.append(line.strip())
                else:
                    id_li.append(line_list[0])
                    if ''.join(line_list[1:6]) == '-----':continue
                    if line_list[0]==sampleid:
                        filter_li.append(line.strip())
                    else:
                        pass
        if sampleid not in id_li:
            sys.stderr.write('sample %s is not in %s'%(sampleid,filterFile))
            sys.exit(1)
    except Exception as e:
        traceback.print_exc(e)
    return filter_li,filter_fu

def prepare_allDic(files,disease_file,filesfu):
    disease_Dic = defaultdict(list)
    site_Dic = defaultdict(list)
    site_li = []
    gene_li= []
    key_li=[]
    het_hom ={'het':u'杂合','hom':u'纯合','unknown':u'未明','hem':u'半合子'} #wph add hem 181212
    data_Dic=read_GXYdata(disease_file)
    try:
        for line in files:
            line_list=[x.strip() for x in line.strip().split('\t')]
            if len(line_list) < 89:
                sys.stderr.write('filter文件列数不足，请检查')
                sys.exit(1)
            ##获取disease_Dic
            gene=line_list[85].strip()
            subtype_disease=re.sub(' ','',line_list[83].strip())
            key=gene+':'+subtype_disease
            if  key not in key_li:
                key_li.append(key)
                if key in data_Dic:
                    disease_detail=data_Dic[key]
                    biaoxian=line_list[89]
                    genetic_mode=line_list[84]
                    dalei_miaoshu=line_list[79]
                    zhibingxing=line_list[78]
                    zhiliao=line_list[88]
                    disease_detail.append(genetic_mode)
                    disease_detail.append(dalei_miaoshu)
                    disease_detail.append(biaoxian)
                    disease_detail.append(zhibingxing)
                    disease_detail.append(zhiliao)
                    disease_Dic[gene].append(disease_detail)  #key>gene：subtype_disease,value>[大类+亚类，亚类，大类，遗传模式，疾病描述，临床表现,致病性,治疗建议]
                else:
                    sys.stderr.write(u'请检查filter内%s与GXY.disease.table中的gene名、疾病亚型名是否一致'%key)#gene_dis_key2)
                    sys.exit(1)
            else:
                pass
            ##获取site_Dic
            site=line_list[1]+":"+line_list[2]+":"+line_list[3]+":"+line_list[4]+":"+line_list[5] #chr:start:end:ref:alt
            if site not in site_li:
                if line_list[90]=='-' or line_list[90]=='.' or line_list[90]=='.//.,':    ### panqi 180817
                    exon='-'
                else:
                    exons=line_list[90].split(':')
                    exon=exons[0]+'\n'+exons[1]
                chr_site=get_chr(line_list[1],line_list[2],line_list[3])
                mute_type=get_mut_type(line_list[9])
                level_value=get_level_value(line_list[78])
                site_detail=[line_list[86],line_list[87],chr_site,exon,het_hom[line_list[31].lower()],mute_type,line_list[78],line_list[82],line_list[80],line_list[81],line_list[88],level_value,line_list[44],line_list[83],line_list[84]]
                #nue,     acid,    chr_site,exon,het/hom,                       mute_type,zhibingxing,    literature,  gene_des,     site_des,      advise,      level_value   true_false    subtype       genetic_model
                site_Dic[line_list[85]].append(site_detail)
                site_li.append(site)
            else:
                pass
    except Exception as e:
        traceback.print_exc(e)
    site_Dic_fubiao = defaultdict(list)
    site_li_fubiao = []
    try:
        for line in filesfu:
            line_list=[x.strip() for x in line.strip().split('\t')]
            site=line_list[1]+":"+line_list[2]+":"+line_list[3]+":"+line_list[4]+":"+line_list[5]+line_list[83]
            if site not in site_li_fubiao:
                if line_list[90]=='-' or line_list[90]=='.' or line_list[90]=='.//.,':
                    exon='-'
                else:
                    exons=line_list[90].split(':')
                    exon=exons[0]+'\n'+exons[1]
                mute_type=get_mut_type(line_list[9])
                level_value=get_level_value(line_list[78])
                site_detail=[line_list[86],line_list[87],exon,het_hom[line_list[31].lower()],mute_type,line_list[77],line_list[78],line_list[83],line_list[84],line_list[82],level_value]
                site_Dic_fubiao[line_list[85]].append(site_detail)
                site_li_fubiao.append(site)
    except Exception as e:
        traceback.print_exc(e)
    sort_site = sort_siteDic(site_Dic)
    sort_site_fubiao = sort_siteDic_fubiao(site_Dic_fubiao)
    return disease_Dic,sort_site,sort_site_fubiao

def sort_siteDic(Dic):
    sorted_Dic={}
    for key in Dic:
        sorted_Dic[key]=sorted(Dic[key],key=lambda x:x[11])  #### panqi x[11]
    return sorted_Dic

def sort_siteDic_fubiao(Dic):
    sorted_Dic={}
    for key in Dic:
        sorted_Dic[key]=sorted(Dic[key],key=lambda x:x[10])
    return sorted_Dic

def get_chr(Chr,start,end):
    site=''
    if start=='.' or start=='-' or Chr=='.' or Chr=='-':
        sys.stderr.write(u'染色体号或位点起始为空')
        sys.exit(1)
    if end=='.' or end=='-' or start == end:
        site='Chr'+Chr+':'+start
    else:
        site='Chr'+Chr+':'+start+' - '+end
    return site

def get_level_value(level):
    level_value=0
    if level=='致病':level_value=1
    elif level=='可能致病':level_value=2
    elif '临床意义未明1级' in re.sub(' ','',level): level_value=3
    elif '临床意义未明2级' in re.sub(' ','',level): level_value=4
    elif '临床意义未明3级' in re.sub(' ','',level): level_value=5
    elif '临床意义未明4级' in re.sub(' ','',level): level_value=6
    elif '临床意义未明5级' in re.sub(' ','',level): level_value=7
    elif level=='可能良性': level_value=8
    elif level=='良性':level_value=9
    else:
        sys.stderr.write(u'请检查致病等级')
        sys.exit(1)
    return level_value

def get_mut_type(mute):
    mute_Dic={'frameshift deletion':'移码缺失变异',
    'frameshift insertion':'移码插入变异',
    'frameshift substitution':'移码替换变异',
    'nonframeshift deletion':'非移码缺失变异',
    'nonframeshift insertion':'非移码插入变异',
    'nonframeshift substitution':'非移码替换变异',
    'stopgain':'无义变异',
    'stopgain2':'缺失移码变异',
    'stoploss':'终止缺失变异',
    'stoploss SNV':'终止缺失变异',
    'synonymous SNV':'同义变异',
    'nonsynonymous SNV':'错义变异',
    'SNV':'单核苷酸变异',
    'splicing SNV':'剪切区单核苷酸变异',
    }
    if mute in mute_Dic:
        mute_des=mute_Dic[mute]
    else:
        sys.stderr.write(u'请检查filter文件第9列是否为%s'%mute_Dic.keys())
        sys.exit(1)
    return mute_des

def read_GXYdata(disease_file):
    data=disease_file
    disease_Dic=defaultdict(list)
    if os.path.isfile(data) == False:
        sys.stderr.write('%s is not a file' % data)
        sys.exit(1)
    try:
        with open(data, 'rU') as file_handle:
            for line in file_handle:
                line_list=[x.strip() for x in line.strip().split('\t')]
                if line_list[0]==u'疾病（共15种）（疾病列表中）':continue
                if len(line_list) < 4:
                    sys.stderr.write(u'cs.data 列数不足')
                    sys.exit(1)
                gene=re.sub(' ','',line_list[1])
                subtype_disease=re.sub(' ','',line_list[2])
                key=gene+':'+subtype_disease
                disease_Dic[key]=[line_list[0],line_list[2],line_list[3]]  # 大类+亚型，亚型名，大类名  
    except Exception as e:
        traceback.print_exc(e)
    return disease_Dic
	
def get_zongshu(site_Dic):
    cover_Dic={}
    hom_gene,hom_level,het_gene,het_level = tongji_hethom(site_Dic)
    hom_gene_set=[x for x in set(hom_gene)]
    het_gene_set=[x for x in set(het_gene)]
    if len(het_gene)==0:    ##全为纯合突变
        if len(hom_gene)==1:
            zongshu='本次受检样本中共检出 1 个纯合变异，存在于%s基因上，'%(hom_gene[0])
        elif len(set(hom_gene))==1:
            zongshu='本次受检样本中共检出 %s 个纯合变异，均存在于%s基因上，'%(str(len(hom_gene)),hom_gene_set[0])
        else:
            gene='、'.join(hom_gene)
            zongshu='本次受检样本中共检出 %s 个纯合变异，分别存在于%s基因上，'%(str(len(hom_gene)),gene)
    elif len(hom_gene)==0:  ##全为杂合突变
        if len(het_gene)==1:
            zongshu='本次受检样本中共检出 1 个杂合变异，存在于%s基因上，'%(het_gene[0])
        elif len(set(het_gene))==1:
            zongshu='本次受检样本中共检出 %s 个杂合变异，均存在于%s基因上，'%(str(len(het_gene)),het_gene_set[0])
        else:
            gene='、'.join(het_gene)
            zongshu='本次受检样本中共检出 %s 个杂合变异，分别存在于%s基因上，'%(str(len(het_gene)),gene)
    else:
        zongshu='本次受检样本中共检出 %s 个变异，%s个杂合变异和%s个纯合变异，分别存在于%s和%s基因上，'%(str(len(het_gene)+len(hom_gene)),str(len(het_gene)),str(len(hom_gene)),'、'.join(het_gene),'、'.join(hom_gene))
    zongshu += get_level(hom_level,het_level)
    zongshu += '与上述基因相关的疾病及检出位点详细信息请见下表。'
    cover_Dic['detect_summary'] = zongshu
    return cover_Dic

def tongji_hethom(site_Dic):
    hom_gene = []
    het_gene = []
    hom_level = []
    het_level = []
    for key in sorted(site_Dic,key=lambda x:site_Dic[x][0][11]):
        value=site_Dic[key]
        for line in value:
            if line[4] == "纯合":
                hom_gene.append(key)
                hom_level.append(line[6])
            else:
                het_gene.append(key)
                het_level.append(line[6])
    return hom_gene,hom_level,het_gene,het_level

def get_level(hom,het):
    zong=hom+het
    Filt=set(zong)
    filt=[i for i in Filt]
    miaoshu=''
    if len(zong)==1:
        miaoshu = '严格遵循 ACMG 解读标准，判定为%s。'%zong[0]
    elif len(filt)==1:
        miaoshu = '严格遵循 ACMG 解读标准，均判定为%s。'%filt[0]
    elif len(het)==0:
        hom_miaoshu='、'.join(hom)
        miaoshu = '严格遵循 ACMG 解读标准，依次判定为%s。'%hom_miaoshu
    elif len(hom)==0:
        het_miaoshu='、'.join(het)
        miaoshu = '严格遵循 ACMG 解读标准，依次判定为%s。'%het_miaoshu
    else:
        hom_miaoshu='、'.join(hom)
        het_miaoshu='、'.join(het)
        miaoshu = '严格遵循 ACMG 解读标准，依次判定为%s和%s。'%(het_miaoshu,hom_miaoshu)
    return miaoshu

def get_disease_table(Dic,siteDic):
    #疾病+亚型0，疾病亚型1，疾病大类2，遗传模式3，疾病描述4，临床表现5，致病性6,指导建议7
    #nue0,acid1,chr_site2,exon3,het/hom4,mute_type5,
    #zhibingxing6,literature7,gene_des8,site_des9,advise10,level_value11,true_false12
    cover_Dic={}
    tabledis = []
    Dic_new = Dic
    ADAR=[]
    table=''
    sortDic=siteDic
    for key in sorted(sortDic, key=lambda x: sortDic[x][0][11]):  ##根据致病性排序
        gene_line = []

        for line in Dic_new[key]:
            ADAR.append(line[3])
            lin = {'n1': line[0], 'n2': line[3], 'n3': line[5]}
            gene_line.append(lin)
        keylin = {'n0': key, 'area': gene_line}
        tabledis.append(keylin)
    """
    for key in sorted(sortDic,key=lambda x:sortDic[x][0][11]):
        if len(Dic[key])==1:
            ADAR.append(Dic[key][0][3])
            table += r'''
%s & \makecell[c]{%s} & %s & %s \\ \hline'''%(key,Dic[key][0][0],Dic[key][0][3],Dic[key][0][5])
        else:
            n=0
            for li in Dic[key]:
                ADAR.append(li[3])
                n += 1
                if n==1:
                    table += r'''
\multirow{%s}{*}{%s} & \makecell[c]{%s} & %s & %s \\ \cline{2-4}'''%(str(len(Dic[key])),key,li[0],li[3],li[5])
                elif n==len(Dic[key]):
                    table += r'''
& \makecell[c]{%s} & %s & %s \\ \hline'''%(li[0],li[3],li[5])
                else:
                    table += r''' 
& \makecell[c]{%s} & %s & %s \\ \cline{2-4}'''%(li[0],li[3],li[5])
    genetic_model = get_model(set(ADAR))
    """
    cover_Dic['tbl_contents']=tabledis
    cover_Dic['genetic_model'] = get_model(set(ADAR))
    #cover_Dic['result_table']=table
    #cover_Dic['genetic_model']=genetic_model
    return cover_Dic

def prepare_lite(li):
    li_new=[]
    for i in li:
        if '|' in i:
            arr=i.split('|')
            for k in arr:
                if k=='' or k=='-' or k=='.':
                    pass
                else:
                    li_new.append(k)
        else:
            li_new.append(i)
    return li_new

def get_lite(lia,lib):
    lite=[]
    liteanum=''
    litebnum=''
    lia_New = prepare_lite(lia)
    lib_New = prepare_lite(lib)
    lia_new = [i for i in set(lia_New)]
    lib_new = [i for i in set(lib_New)]
    li=lia_new + lib_new
    ##获取文献列表
    if len(li)==0:
        pass
    else:
        for index,item in enumerate(li):
            lite.append(item)
    ##获取索引数字
    suoyin=[]
    if len(lia_new)==0:
        liteanum=''
    else:
        for index,item in enumerate(lia_new):
            suoyin.append(str(index+1))
        liteanum = ','.join(suoyin)
        liteanum = '[' + liteanum +']'
    suoyin=[]
    if len(lib_new)==0:
        litebnum=''
    else:
        for index,item in enumerate(lib_new):
            suoyin.append(str(index+1+len(lia_new)))
        litebnum = ','.join(suoyin)
        litebnum = '[' + litebnum +']'
    return liteanum,litebnum,lite

def get_model(li):
    genetic_model=''
    mode_Dic={'AD':'AD：常染色体显性遗传',
    'AR':'AR：常染色体隐性遗传',
    'XLD':'XLD：X 染色体连锁显性遗传',
    'XLR':'XLR：X 染色体连锁隐性遗传',
    'XL':'XL：X 染色体连锁遗传',
    'AD/AR':'AD/AR：常染色体遗传',
    'AR/AD':'AR/AD：常染色体遗传',
    'AD,AR':'AD/AR：常染色体遗传',
    'AR,AD':'AR/AD：常染色体遗传',
    'AD，AR':'AD/AR：常染色体遗传',
    'AR，AD':'AR/AD：常染色体遗传',
    'AD，Smu':'AD：常染色体显性遗传；Smu：体细胞突变'
    }
    for mode in li:
        if mode in mode_Dic:
            genetic_model+=mode_Dic[mode]
        elif mode == r'不明':
            pass
        else:
            sys.stderr.write(u'遗传模式为%s，不在AD AR XLD XLR XL AD/AR AR/AD内' % mode)
            sys.exit(1)
    """
    ad_f=0
    for i in li:
        mode=re.sub(' ','',i)
	if mode=='AD':
	    if ad_f==0:
		ad_f=1
		genetic_model.append( mode_Dic[mode])
	elif mode == 'AD，Smu':
	    if ad_f==0:
		ad_f=1
		genetic_model.append( mode_Dic[mode])
	    else:
		genetic_model.append('Smu：体细胞突变')
        elif mode_Dic.has_key(mode):
            genetic_model.append( mode_Dic[mode])
        else:
            sys.stderr.write(u'遗传模式为%s，不在AD AR XLD XLR XL AD/AR AR/AD AD,AR AR,AD内'%mode)
            sys.exit(1)
    genetic='；'.join(genetic_model) 
        """

    return genetic_model

def get_site_table(Dic):
    table=[]
    pep=''
    cover_Dic={}
    num=0
    sortDic=Dic

    """
    for key in sorted(sortDic,key=lambda x:sortDic[x][0][11]):
        if len(sortDic[key])==1:
            color=get_color(sortDic[key][0][6])
            pep = sortDic[key][0][1]
            if "fs" in sortDic[key][0][1]:
                pep = "\makecell*[c]{%s}" % sortDic[key][0][1].replace("fs","\\\\fs") #wph add 190612
            table += r'''
\color{MyFontGray}
%s  & %s & %s & %s &  \makecell*[c]{%s} & %s & %s & \textcolor{%s}{%s} \\ \hline
'''%(key,sortDic[key][0][0],pep,sortDic[key][0][2],sortDic[key][0][3],sortDic[key][0][4],sortDic[key][0][5],color,sortDic[key][0][6])
        else:
            num=get_multline(sortDic[key])
            m=0
            for line in sortDic[key]:
                m += 1
                color=get_color(line[6])
                pep=line[1]
                if "fs" in line[1]:
                    pep= "\makecell*[c]{%s}" % line[1].replace("fs","\\\\fs")  #wph add 190612
                if m==1:
                    table += r'''
\color{MyFontGray}
\multirow{%s}{*}{%s} & %s & %s & %s &  \makecell*[c]{%s} & %s & %s & \textcolor{%s}{%s} \\ \cline{2-8}'''%(str(num),key,line[0],pep,line[2],line[3],line[4],line[5],color,line[6])
                elif m==len(Dic[key]):
                    table += r'''
\color{MyFontGray}
& %s & %s & %s &  \makecell*[c]{%s} &  %s & %s & \textcolor{%s}{%s} \\ \hline'''%(line[0],pep,line[2],line[3],line[4],line[5],color,line[6])
                else:
                    table += r'''
\color{MyFontGray}
& %s & %s & %s &  \makecell*[c]{%s} &  %s & %s & \textcolor{%s}{%s} \\ \cline{2-8}'''%(line[0],pep,line[2],line[3],line[4],line[5],color,line[6])
    cover_Dic['site_table']=table
    """
    for key in sorted(sortDic, key=lambda x: sortDic[x][0][11]):

        num = get_multline(sortDic[key])

        for line in sortDic[key]:
            aa = line[6]
            if aa == '致病':
                aa = RichText('致病', bold=True)
            lis = {'n0': key, 'n1': line[0], 'n2': line[1], 'n3': line[2], 'n4': RichText(line[3]), 'n5': line[4], 'n6': line[5],
                   'n7': aa} #RichText 使用换行
            table.append(lis)
    cover_Dic['tbl_contents2'] = table

    return cover_Dic

def get_fubiao_table(siteDic):
    cover_Dic={}
    table=[]
    pep=''
    sortDic=siteDic
    biaoshi=0
    Dic_wenxan=[]
    DicK = []
    for key in sorted(sortDic, key=lambda x: sortDic[x][0][10]):
        for line in sortDic[key]:
            gene_line = []
            lin = {"n1": sortDic[key][0][7], "n2": sortDic[key][0][8]}
            gene_line.append(lin)
            keylin = {'n0': key, 'n1': line[0], 'n2': line[1], 'n3': RichText(line[2]), 'n4': line[3], 'n5': line[4],
                      'n6': line[5], 'n7': line[6], 'area': gene_line}
            table.append(keylin)


        DicK.append(sortDic[key][0][9])

    wenxian = get_wenxianfu(list(set(DicK)))
    """
    for key in sorted(sortDic,key=lambda x:sortDic[x][0][10]):
    #for key in  sortDic:
        biaoshi=1
        m1=len(sortDic[key])
        if len(sortDic[key])==1:
            pep = sortDic[key][0][1]
            if "fs" in sortDic[key][0][1]:
                pep = "\makecell*[c]{%s}" % sortDic[key][0][1].replace("fs","\\\\fs")
            table += r'''
\color{MyFontGray}
%s  & %s & %s &  \makecell*[c]{%s} & %s & %s & %s & \makecell*[c]{%s} & %s  & %s  \\ \hline
'''%(key,sortDic[key][0][0],pep,sortDic[key][0][2],sortDic[key][0][3],sortDic[key][0][4],sortDic[key][0][5],sortDic[key][0][6].replace("临床意义","临床意义\\\\"),sortDic[key][0][7],sortDic[key][0][8])
            if sortDic[key][0][9] not in Dic_wenxan:
                Dic_wenxan.append(sortDic[key][0][9]) 
        else:
            m=0;
            num=len(sortDic[key])
            for line in sortDic[key]:
                m += 1
                pep=line[1]
                if "fs" in line[1]:
                    pep= "\makecell*[c]{%s}" % line[1].replace("fs","\\\\fs")
                if  m==1:
                    table += r'''
\color{MyFontGray}
\multirow{%s}{*}{%s}  & %s & %s &  \makecell*[c]{%s} & %s & %s & %s & \makecell*[c]{%s} & %s  & %s  \\  \cline{2-10}
'''%(str(num),key,line[0],pep,line[2],line[3],line[4],line[5],line[6].replace("临床意义","临床意义\\\\"),line[7],line[8])
                elif m==num:
                    table += r'''
\color{MyFontGray}
  & %s & %s &  \makecell*[c]{%s} & %s & %s & %s & \makecell*[c]{%s} & %s  & %s  \\  \hline
'''%(line[0],pep,line[2],line[3],line[4],line[5],line[6].replace("临床意义","临床意义\\\\"),line[7],line[8])
                else:
                    table += r'''
\color{MyFontGray}
  & %s & %s &  \makecell*[c]{%s} & %s & %s & %s & \makecell*[c]{%s} & %s  & %s  \\  \cline{2-10}
'''%(line[0],pep,line[2],line[3],line[4],line[5],line[6].replace("临床意义","临床意义\\\\"),line[7],line[8])
                if line[9] not in Dic_wenxan:
                    Dic_wenxan.append(line[9])    
    biaotou=r'''\newpage
\includegraphics[height=0.4cm]{XKmini.pdf} \zihao{4}{\color{DarkBlue} \sym{ 附表1：}}\par
\par \vspace*{6mm}
\begin{spacing}{1.3}
\color{MyFontGray} \zihao{5}
下表为无ACMG证据、证据矛盾（偏致病类证据与偏良性证据共存）或偏良性证据的位点（可能与相关疾病出现的表型有关，具体情况请结合临床综合分析）
\par \vspace*{6mm}
\arrayrulecolor{MyFontGray}
\renewcommand\arraystretch{1.4}
\fontsize{7}{8.4}\selectfont \color{MyFontGray}
\tablefirsthead{
\rowcolor{DarkBlue}  
\color{white}{\sym{基因}} & \color{white}{\sym{ 核苷酸变异}} & \color{white}{\sym{氨基酸变异}}  &  \makecell[c]{\color{white} \sym{转录本}\\ \sym{外显子编号}}  & \color{white}{\sym{变异状态}} & \color{white}{\sym{变异类型}} &  \makecell[c]{\color{white} \sym{ACMG} \\\sym{条目}}  &  \color{white}{\sym{位点致病性}} & \color{white}{\sym{相关疾病}} & \color{white}{\sym{遗传模式}} \\}
\tablehead{}
\tabletail{\hline}
\tablelasttail{}
\begin{supertabular}{| m{9mm}<{\centering} | m{15mm}<{\centering} | m{13mm}<{\centering} | m{15mm}<{\centering} | m{6mm}<{\centering} | m{10mm}<{\centering} | m{10mm}<{\centering} |m{13mm}<{\centering}| m{31mm}<{\centering} | m{6mm}<{\centering} | }
'''
    biaowei=r'''\end{supertabular}
%\par \vspace*{3mm}
%\zihao{6}
%注：关于频率，“.”表示在对应数据库中未收录该位点
\end{spacing}'''
    fubiao_tablea=''
    fubiao_wenxiana=''
    if biaoshi==1:
        fubiao_tablea=biaotou+table+biaowei
        fubiao_wenxiana=get_wenxianfu(Dic_wenxan)
    """
    cover_Dic['wenxian'] = wenxian
    cover_Dic['tbl_contents3'] = table
    return cover_Dic

def get_multline(li):
    level=[]
    num=0
    for line in li:
        if u'临床意义未明' in line[6]:
            num += 2
        else:
            num += 1
    return num

def get_color(cli):
    color=''
    if cli ==u"致病": color="red"
    elif cli ==u"可能致病": color="orange"
    elif cli ==u"良性" or cli ==u"可能良性": color="MyGreen"
    else: color="yellow"
    return color

def get_jiexi_table(disease,site):
    #nue0,acid1,chr_site2,exon3,het/hom4,mute_type5,
    #zhibingxing6,literature7,gene_des8,site_des9,advise10,level_value11
    site_sortDic=site
    disease_Dic=disease
    cover_Dic={}
    sd = tpl.new_subdoc()
    count = 0
    litea = []
    for key in sorted(site_sortDic, key=lambda x: site_sortDic[x][0][10]):
        if count > 0:
            sd.add_page_break()
        count = count + 1
        jiexi1 = ''
        jiexi2 = ''
        jiexi3 = ''
        disease_li, disease_des, site_li, site_des, level = prepare_jiexi(site_sortDic[key], disease_Dic[key])
        sublite = get_wenxian(site_sortDic[key])

        for item in sublite:
            litea.append(item)

        if len(disease_li) != len(disease_des):
            sys.stderr.write(u'请检查filter是否每个位点都包含疾病名称、疾病描述')
            sys.exit(1)

        if len(site_li) != len(level):
            sys.stderr.write(u'位点数与致病等级数不一致')
            sys.exit(1)

        # jiexi2 = r''' %s '''%site_des[0]

        jiexi3 = r'''  %s ''' % site[key][0][8]

        row_dis = len(disease_li)

        table2 = sd.add_table(1, 2, style=sd.styles['newstyle-1'])
        table2.style.font.size = Pt(9)
        table2.style.font.name = u'微软雅黑'
        col = table2.columns[0]
        col.width = Inches(2.5)
        col = table2.columns[1]
        col.width = Inches(9)
        table2.cell(0, 0).merge(table2.cell(0, 1))
        row_cells = table2.rows[0].cells

        if len(site_li) == 1:
            p = row_cells[0].paragraphs[0]
            print (site_li[0])
            run = p.add_run(u'风险位点: ' + key + '——' + str(site_li[0]))

            p.paragraph_format.first_line_indent = Inches(0.2)
            run.font.size = Pt(9)
            p = row_cells[0].add_paragraph()
            run = p.add_run(text=u'相关疾病: ' + u'、'.join(disease_li))
            p.paragraph_format.first_line_indent = Inches(0.2)
            # run.bold = True
            run.font.size = Pt(9)
            p = row_cells[0].add_paragraph()
            run = p.add_run(u'位点致病性: ' + level[0])
            p.paragraph_format.first_line_indent = Inches(0.2)
            run.font.size = Pt(9)

        else:
            for index, value in enumerate(site_li):
                if index == 0:
                    p = row_cells[0].paragraphs[0]
                else:
                    p = row_cells[0].add_paragraph()
                run = p.add_run(u'风险位点' + str(index + 1) + ': ' + key + '——' + site_li[index] + ', ' + level[index])
                p.paragraph_format.first_line_indent = Inches(0.2)
                run.font.size = Pt(9)

            p = row_cells[0].add_paragraph()
            run = p.add_run(text=u'相关疾病: ' + u'、'.join(disease_li))
            p.paragraph_format.first_line_indent = Inches(0.2)
            run.font.size = Pt(9)

        if len(site_des) == 1:
            row_cells = table2.add_row().cells
            p = row_cells[0].paragraphs[0]
            run = p.add_run(u'变异解析')
            run.font.size = Pt(9)
            row_cells[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            jiexi2 = r''' %s ''' % site_des[0]
            p = row_cells[1].paragraphs[0]
            run = p.add_run(jiexi2)
            p.paragraph_format.first_line_indent = Inches(0.2)
            run.font.size = Pt(9)
        elif len(site_des) > 1:
            for index, value in enumerate(site_des):
                row_cells = table2.add_row().cells
                p = row_cells[0].paragraphs[0]
                run = p.add_run(u'变异解析' + str(index + 1))
                run.font.size = Pt(9)
                row_cells[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                jiexi2 = r''' %s ''' % site_des[index]
                p = row_cells[1].paragraphs[0]
                run = p.add_run(jiexi2)
                p.paragraph_format.first_line_indent = Inches(0.2)
                run.font.size = Pt(9)

        row_cells = table2.add_row().cells
        p = row_cells[0].paragraphs[0]
        run = p.add_run(u'基因描述')
        run.font.size = Pt(9)
        row_cells[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
        row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        p = row_cells[1].paragraphs[0]
        run = p.add_run(jiexi3)
        p.paragraph_format.first_line_indent = Inches(0.2)
        run.font.size = Pt(9)

        for index, item in enumerate(disease_li):
            row_cells = table2.add_row().cells
            # row_cells[0].text = unicode(item,'utf-8')
            p = row_cells[0].paragraphs[0]
            run = p.add_run(item)
            run.font.size = Pt(9)
            row_cells[0].paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
            row_cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = row_cells[1].paragraphs[0]
            run = p.add_run(disease_des[index])
            p.paragraph_format.first_line_indent = Inches(0.2)
            run.font.size = Pt(9)
        sd.add_paragraph("")
    jiexi_beizhu = r'''
    预测软件包括SIFT、PolyPhen2、M-CAP 和REVEL, 用于预测错义变异是否会导致蛋白结构和功能发生改变, 准确率大致在65%-80%。需要说明的是, 各软件提供的只是基于算法的预测结果, 可以为位点解读提供参考, 但不足以作为致病性判定的唯一标准, 还需要结合其他信息综合判断。以上位点变异为基因检测的客观结果，但变异是否真正致病需要由医生结合临床信息综合判断。如果您想进一步了解基因变异的信息或家族遗传史 ，建议联系家人一起进行检测。'''
    p = sd.add_paragraph()
    run = p.add_run(u'注:')
    run.font.size = Pt(8)
    run.bold = True
    run.font.name = u'微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    p = sd.add_paragraph()
    run = p.add_run(jiexi_beizhu)
    run.font.size = Pt(8)
    run.font.name = u'微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')

    cover_Dic['jiexi_site'] = sd
    cover_Dic['wenxian'] = litea
    return cover_Dic

def get_advise(li):
    advise=r'\par \vspace*{5mm}'
    for i in li:
        if i[6]==u'致病' or i[6]==u'可能致病':
            advise += r'''
\includegraphics[height=0.4cm]{XKmini.pdf} \zihao{4}{\color{DarkBlue}\sye \sym{ %s治疗指导建议}}\par
\fontsize{10.5}{13.65}\selectfont{\color{MyFontGray} %s}\par'''%(i[2],i[7])
    return advise

def get_wenxian_old(li):
    wenxian=[]
    lite=''
    for i in li:
        if i[7]=='' or i[7] =='-':
            pass
        else:
            wenxian.append(i[7])
    wenxian_new=prepare_lite(wenxian)
    if len(wenxian_new) ==0:
        lite=r'''
%\par \vspace*{5mm}'''
    else:
        lite =r'''
\par \vspace*{3mm}
\zihao{6}{\color{DarkBlue} 参考文献}\\ \vspace*{-6mm}
\begin{spacing}{1.5}
\zihao{-6} \color{MyFontGray}
\setdefaultleftmargin{2em}{}{}{}{.5em}{.5em}
\begin{compactenum}'''
        for i in wenxian_new:
            lite +=r'''
\item %s'''%i
        lite += r'''
\end{compactenum}
\end{spacing}'''
    return lite

def get_wenxian(li):
    wenxian=[]
    wenxian_new=[]
    for i in li:
        if i[7]=='' or i[7] =='-':
            pass
        else:
            wenxian.append(i[7])
    wenxian_new=prepare_lite(wenxian)
    return wenxian_new

def get_wenxianfu(li):
    wenxian=[]
    wenxian_new=[]
    for i in li:
        for j in i.split("|"):
            if j=='' or j =='-' or j =='.':
                pass
            else:
                wenxian.append(j)
    wenxian_new=prepare_lite(wenxian)
    return wenxian_new

def get_wenxianfu_old(li):
    wenxian=[]
    lite=''
    for i in li:
            if i=='' or i=='-' or i=='.':
                pass
            else:
                wenxian.append(i)
    wenxian_new=prepare_lite(wenxian)
    if len(wenxian_new) ==0:
        lite=r'''
\par \vspace*{5mm}'''
    else:
        lite =r'''
\par \vspace*{5mm}
\zihao{6}{\color{DarkBlue} 参考文献}\\ \vspace*{-6mm}
\begin{spacing}{1.5}
\zihao{-6} \color{MyFontGray}
\setdefaultleftmargin{2em}{}{}{}{.5em}{.5em}
\begin{compactenum}'''
        for i in set(wenxian_new):
            lite +=r'''
\item %s'''%i
        lite += r'''
\end{compactenum}
\end{spacing}'''
    return lite

def get_head(key,site,disease,level):
    head=''
    disease_new=[]
    for i in disease:
        if i[2]==i[1]:type_sub=i[2]
        else:type_sub=i[2]+'（'+i[1]+'）'
        disease_new.append(type_sub)
    if len(site) != len(level):
        sys.stderr.write(u'位点数与致病等级数不一致')
        sys.exit(1)
    else:
        head =r'''
\begin{spacing}{1.2}
\zihao{5} \color{white} \sym \sye
\renewcommand\arraystretch{1}
\colorbox{DarkBlue}{
\begin{tabular}{L{2cm} L{14cm}}'''
        for index,item in enumerate(site):
            head +=r'''
风险位点%s: & %s—— %s, %s\\'''%(str(index+1),key,item,level[index])
        head +=r'''
相关疾病: & %s\\
\end{tabular}}
\end{spacing} '''%('、'.join(disease_new))
    return head

def get_disease_row(disease,disease_des):
    table=''
    if len(disease) != len(disease_des):
        sys.stderr.write(u'请检查filter是否每个位点都包含疾病名称、疾病描述')
        sys.exit(1)
    else:
        for index,item in enumerate(disease):
            table += r'''
%s & %s \\'''%(item.replace("\\",""),disease_des[index])
    return table

def prepare_jiexi(site_li,disease_li): #wph changed 190612
    #nue,acid,chr_site,exon,het/hom/,mute_type,zhibingxing,literature,gene_des,site_des
    disease=[]
    disease_des=[]
    site=[]
    site_des=[]
    level=[]
    for line in disease_li:
        if line[0] not in disease:
            disease.append(line[0])
            disease_des.append(line[4])
        else:
            pass
    for line in site_li:
        site_detail=line[0]+':'+line[1]+', '+line[4]
        site.append(site_detail)
        site_des.append(line[9])
        level.append(line[6])
    return disease,disease_des,site,site_des,level

def get_QC_Inf(QCfile,sampleid):
    QC_inf_list=[]
    QC_li=[]
    cover_Dic= defaultdict(list)
    head=['cover_one','cover_four','cover_ten','cover_20','depave']
    if os.path.isfile(QCfile) == False:
        print('%s is not a file' % QCfile)
        return cover_Dic
    try:
        file_handle = open(QCfile, 'rU')
        QC_list = file_handle.readlines()
        for line in QC_list:
            QC_arr = [x.strip() for x in line.strip().split('\t')]
            if QC_arr[0] == sampleid and len(QC_arr) > 40:
                QC_li.append(QC_arr)
            else:
                pass
        if len(QC_li)==1:
            QC_inf = [QC_li[0][2],QC_li[0][34],QC_li[0][33],QC_li[0][32],QC_li[0][3]]
        else:
            sys.stderr.write(u'请检查质控结果是否无或未合并')
            sys.exit(1)
        cover_Dic = {k:v for k,v in zip(head,QC_inf)}
    except Exception as e:
        traceback.print_exc(e)
    finally:
        if 'fileHandle' in dir() or file_handle.closed == False:
            file_handle.close()
    return cover_Dic

def get_HTB(sampleid):
    # files=glob.glob(r'/PUBLIC/pipline/database/sheet_hutong/*心康互通表*.xls*')
    files = glob.glob(r'/PUBLIC/pipline/database/sheet_hutong/2019心康送样信息单*.xls*')
    max_num = 0
    age = date = ''
    sex = ''
    age_num = ''  # wph changed 181106
    date_num = ''
    for i in files:
        hutong = i.split('/')[-1]
        # st=re.match(r'心康互通表（([0-9]*)）.xlsx',hutong)
        st = re.match(r'2019心康送样信息单-([0-9]*).xlsx', hutong)
        if int(st.group(1)) > max_num:
            max_num = int(st.group(1))
    HTB = r'/PUBLIC/pipline/database/sheet_hutong/2019心康送样信息单-%s.xlsx' % max_num
    try:
        data = xlrd.open_workbook(HTB)
        table = data.sheet_by_name('下单表')
        nrows = table.nrows
        for rownum in range(1, nrows):
            if table.row(rownum)[2].value == sampleid:
                age = table.row(rownum)[8].value
                sex = table.row(rownum)[7].value
                ctype = table.cell(rownum, 1).ctype
                cell = table.cell_value(rownum, 1)
                date
                if ctype == 3:
                    date = datetime(*xldate_as_tuple(cell, 0))
                    date = date.strftime('%Y/%m/%d')
                else:
                    date = table.row(rownum)[1].value
        date_num = date.replace("/", "-")
        if isinstance(age, float) == True:  # wph changed 181106
            age_num = str(int(age)) + "岁"
        elif not age:
            age_num = "不详"
        else:
            age_num = age
        if age_num == '' or date_num == '':
            sys.stderr.write(u"ID不在互通表 或 互通表中没有年龄/送样日期信息")
            exit
    except Exception as e:
        traceback.print_exc(e)
    return age_num, date_num, sex

def get_sampleInf(sampleid, cfgFile, feature_file):
    sample_inf_list = []
    cover_Dic = defaultdict(list)
    head = ['sampleID', 'name', 'sex', 'age', 'data1', 'data2', 'chenghu', 'clinical']
    time2str = time.strftime('%Y-%m-%d', time.localtime(time.time()))
    if os.path.isfile(cfgFile) == False:
        print('%s is not a file' % cfgFile)
        return cover_Dic
    try:
        with open(cfgFile, 'rU') as file_handle:
            for line in file_handle:
                line_list = [x.strip() for x in line.strip().split('\t')]
                if line_list[0] == sampleid:
                    line_list[5] = re.sub(r'\(.*?\)', '', line_list[5])
                    (age, date, sex) = get_HTB(sampleid)  # wph changed 190423 修改性别不匹配的问题
                    if sex not in line_list[5]:
                        line_list[5] = sex
                    chenghu = judge_chenghu(line_list[5])
                    feature = get_feature(feature_file, sampleid)
                    sample_inf_list = [line_list[0], line_list[4], line_list[5], age, date, time2str, chenghu,
                                       feature]  # 需要送样日期，暂时用出报告日期代替
                    break
        cover_Dic = {k: v for k, v in zip(head, sample_inf_list)}
    except Exception as e:
        traceback.print_exc(e)
    return cover_Dic

def judge_chenghu(sex):
    chenghu = ''
    try:
        if sex == r'女':
            chenghu = r'女士'
        elif sex == r'男':
            chenghu = r'先生'
        else:
            print('性别是%s,check please' % sex)
    except Exception as e:
        traceback.print_exc(e)
    return chenghu

def get_feature(feature_file, sampleid):
    feature = '无临床信息'
    if os.path.isfile(feature_file) == False:
        print('%s is not a file' % feature_file)
        return 'Error File'
    try:
        file_handle = open(feature_file, 'rU')
        feature_list = file_handle.readlines()
        for line in feature_list:
            feature_arr = [x.strip() for x in line.strip().split('\t')]
            if feature_arr[0] == sampleid and len(feature_arr) > 1:
                pattern = re.compile(r'临床诊断疾病: ([^\\]+)', re.I)
                m = pattern.match(feature_arr[1])
                if m is None:
                    feature = '-'
                else:
                    feature = r''' %s ''' % (m.groups(0)[0])
                    feature.strip(',|，')
            else:
                pass
    except Exception as e:
        traceback.print_exc(e)
    return feature

def get_sanger_Inf(siteDic):
    cover_Dic = {}
    sortDic = siteDic
    sitDic = []
    for key in sorted(sortDic, key=lambda x: sortDic[x][0][10]):
        for line in sortDic[key]:
            figname = sampleid + '-' + line[2] + '.jpg'
            figname = re.sub('Chr', '', figname)
            figname = re.sub(':', '-', figname)
            figpath = outDir + '/fig/' + figname
            print (figpath)
            site = u'检测位点：' + key + ':' + line[0] + ',' + line[1]
            lis = {}
            if os.path.exists(figpath):
                lis = {'site': site, 'con': InlineImage(tpl, figpath, width=Mm(150), height=Mm(40)), 'space1': '\n',
                       'space2': '\n'}
            else:
                lis = {'site': site, 'con': '参照序列\n样本序列\n\n\n\n', 'space1': '\n', 'space2': '\n'}
            print (site)
            sitDic.append(lis)
    cover_Dic['varient'] = sitDic
    return cover_Dic

def read_sample_info_PCR(figfile):
    het_hom = {'00': '野生纯合', '01': '变异杂合', '11': '变异纯合'}
    sex = read_sample_sex()
    samples = []
    sitetable = []
    sangerperson = {}
    site_genotype = {}
    sitetables = {}
    genotype_hash = {}
    try:
        with open(figfile, 'rU') as file_handle:
            for line in file_handle:
                line_list = [x.strip() for x in line.strip().split('\t')]
                if line_list[1] == sampleid and line_list[2] not in samples:
                    line_list[0] = re.sub('妈妈', '母亲', line_list[0])
                    line_list[0] = re.sub('爸爸', '父亲', line_list[0])
                    ssex = sex[line_list[2]]
                    if line_list[0] not in sangerperson.keys():
                        sangerperson[line_list[0]] = [
                            [line_list[2], line_list[11], line_list[10], line_list[0], line_list[12], ssex]]
                    else:
                        sangerperson[line_list[0]].append(
                            [line_list[2], line_list[11], line_list[10], line_list[0], line_list[12], ssex])
                    samples.append(line_list[2])
                sitesub = line_list[3] + '_' + line_list[4]
                if line_list[2] == sampleid and line_list[2] != '-' and sitesub not in sitetables.keys():
                    site_type = het_hom[line_list[15]]
                    ssex = sex[line_list[2]]
                    if line_list[3] and "X" and ssex == "男" and site_type == "变异纯合":
                        site_type = "半合子";
                    if site_type not in genotype_hash.keys():
                        genotype_hash[site_type] = 0
                    genotype_hash[site_type] = genotype_hash[site_type] + 1
                    sitetables[sitesub] = line_list[6]  # 赋值为基因名字
                    sitetable.append([line_list[3], line_list[4], line_list[5], line_list[6], line_list[7],
                                      line_list[8]])  ##chr,start,rsnum,gene,cds,pep

                if line_list[1] == sampleid:
                    site_type = het_hom[line_list[15]]
                    if line_list[3] and "X" and ssex == "男" and site_type == "变异纯合":
                        site_type = "半合子";
                    if line_list[2] not in site_genotype.keys():
                        site_genotype[line_list[2]] = {}
                    site_genotype[line_list[2]][sitesub] = site_type

    except Exception as e:
        traceback.print_exc(e)
    return sangerperson, sitetable, site_genotype, genotype_hash, sitetables

def read_sample_sex():
    files = glob.glob(r'/PUBLIC/pipline/database/sheet_hutong/2019心康送样信息单*.xls*')
    max_num = 0
    date_num = ''
    for i in files:
        hutong = i.split('/')[-1]
        st = re.match(r'2019心康送样信息单-([0-9]*).xlsx', hutong)
        if int(st.group(1)) > max_num:
            max_num = int(st.group(1))
    HTB = r'/PUBLIC/pipline/database/sheet_hutong/2019心康送样信息单-%s.xlsx' % max_num
    print (HTB)
    id = sampleid;
    id = re.sub('^N(0+)', '', id)
    print (id)
    sample_sex = {}
    try:
        data = xlrd.open_workbook(HTB)
        table = data.sheet_by_name('下单表')
        nrows = table.nrows
        guanxis = []
        for rownum in range(1, nrows):
            tag = str(table.row(rownum)[6].value)
            tag2 = str(table.row(rownum)[6].value)
            if table.row(rownum)[2].value == sampleid:
                sample_sex[sampleid] = table.row(rownum)[7].value
                print (table.row(rownum)[7].value)
            elif u'一代验证' in tag:
                pattern = re.compile(r'一代验证-(\d+)-(.*)', re.I)
                m = pattern.match(tag2)
                number = ''
                if m is None:
                    number = 'nnn'
                else:
                    number = r''' %s ''' % (m.groups(0)[0])
                    number = re.sub(' ', '', number)
                if number == id:
                    if m.groups(0)[1] is None:
                        sys.stderr.write('%s is not a right sanger type' % table.row(rownum)[6])
                        sys.exit(1)
                    guanxi = r''' %s ''' % (m.groups(0)[1])
                    guanxi = re.sub('妈妈', '母亲', guanxi)
                    guanxi = re.sub('爸爸', '父亲', guanxi)
                    sample_sex[table.row(rownum)[2].value] = table.row(rownum)[7].value
    except Exception as e:
        traceback.print_exc(e)
    return sample_sex

def get_sanger_Inf_Family():
    genotype_hash = ['野生纯合', '变异杂合', '变异纯合', '半合子'];
    fig = outDir + '/fig/' + sampleid + '.family.site.txt';
    cmd='cut -f3 %s |sort|uniq|wc -l' % fig
    flag=os.popen(cmd).read()
    if int(flag)<3:
        return 0
    cover_Dic = {}
    info_line = []
    info_line2 = []
    sitDic = []
    col_label = []
    conclude = ''
    print (fig)
    if os.path.exists(fig):
        (sangerpersoninfo, siteDic, siteGenotype, genotype_hash, sitetables) = read_sample_info_PCR(fig)
        conclude = sangertexconclu(sangerpersoninfo, siteDic, siteGenotype, genotype_hash, sitetables)
        line = sangerpersoninfo[u'先证者'][0]
        lin = {'n1': line[2], 'n2': line[1], 'n3': line[3], 'n4': line[4]}
        info_line.append(lin)
        col_label = [u'所在基因', line[2] + u'(先证者)']
        all_sample = [line[0]]
        cols = []
        relat = {}

        for key in sangerpersoninfo:
            if key != '先证者':
                for line in sangerpersoninfo[key]:
                    lin = {'n1': line[2], 'n2': line[1], 'n3': line[3], 'n4': line[4]}
                    info_line.append(lin)
                    newperson = line[2] + '(' + line[3] + ')'
                    relat[line[0]] = [line[3], line[2]]
                    col_label.append(newperson)
                    all_sample.append(line[0])

        for line in siteDic:
            lables = line[2] + ' ' + '(' + line[4] + ',' + line[5] + ')'
            site_sample = line[0] + '_' + line[1]
            cols = [line[3]]
            for sam in all_sample:
                Genotype = siteGenotype[sam][site_sample]
                cols.append(Genotype)
            lin = {'label': lables, 'cols': cols}
            info_line2.append(lin)

        for sam in all_sample:
            samplesigh = []
            count = 0
            for line in siteDic:
                count = count + 1
                site_sample = line[0] + '_' + line[1]
                siteu = u'检测位点：' + line[3] + ':' + line[4] + ',' + line[5]
                print (siteu)
                if sam != sampleid:
                    figname = sam + '-' + line[0] + '-' + line[1] + '.jpg'
                    figname = re.sub('Chr', '', figname)
                    figname = re.sub(':', '-', figname)
                    figpath = '/XK1/proj01/190618_2012//report/02individual/N004219/fig/' + figname
                    print (figpath)
                    ss = relat[sam][1] + u'(先证者' + relat[sam][0] + ')'
                    if count > 1:
                        ss = ''
                    lis = {}
                    print(ss)
                    if os.path.exists(figpath):
                        lis = {'sample': ss, 'site': siteu,
                               'con': InlineImage(tpl, figpath, width=Mm(150), height=Mm(40)), 'space1': '\n',
                               'space2': '\n'}
                    else:
                        lis = {'sample': ss, 'site': siteu, 'con': '参照序列\n样本序列\n\n\n\n', 'space1': '\n', 'space2': '\n'}

                    sitDic.append(lis)

    cover_Dic['col_labels'] = col_label
    cover_Dic['sangerSam'] = info_line
    cover_Dic['sangerResult'] = info_line2
    cover_Dic['varientFamily'] = sitDic
    cover_Dic['descript'] = conclude

    return cover_Dic

def sangertexconclu(sangerperson, sitetable, siteGenotype, genotype_hash, sitetables):
    conclu = ''
    concluall = ''
    genotyp_count = ''
    geneall = ''
    mutation = ''
    unmutation = ''
    father_flag = 0
    father_ID = ''
    mother_flag = 0
    mother_ID = ''
    count = {}
    family = [];
    genotype_list = genotype_hash
    siteList = sitetables
    geneDic = []
    persen = sangerperson
    siteDic = sitetable
    propositusID = ''
    #    hh = RichText('\n')

    for key in genotype_list:
        if genotype_list[key]:
            genotyp_count = genotyp_count + str(genotype_list[key]) + u'个' + key + '，'
    print(genotyp_count)

    genotyp_count = re.sub('变异杂合', '杂合变异', genotyp_count)
    genotyp_count = re.sub('变异纯合', '纯合变异', genotyp_count)
    genotyp_count = re.sub('半合子', '半合子变异', genotyp_count)

    for key in siteList:
        geneDic.append(siteList[key])

    set(geneDic)

    geneall = '、'.join(geneDic)

    concluall = '先证者样本中检测出 %s 存在于 %s 基因上，变异位点详情如上述。本报告为先证者 %s 的家系验证报告，一代测序结果显示:' % (
    genotyp_count, geneall, persen['先证者'][0][2])

    site_count = len(siteList)

    propositusID = persen['先证者'][0][0]
    for key in persen:
        if key not in family:
            family.append(key)
        if key == "母亲":
            mother_flag = len(persen['母亲'])
            mother_ID = persen['母亲'][0][0]
            if mother_flag > 1:
                sys.stderr.write('%s 母亲数目大于1' % mother_flag)
                sys.exit(1)

        if key == "父亲":
            father_flag = len(persen['父亲'])
            father_ID = persen['父亲'][0][0]
            if father_flag > 1:
                sys.stderr.write('%s 父亲数目大于1' % mother_flag)
                sys.exit(1)

    print ("位点数目" + str(site_count) + "\n")

    for l in siteDic:
        concluall = concluall + RichText('\n')
        concluall = concluall + " %s 基因上的变异（ %s ，%s ）" % (l[3], l[4], l[5])
        sitesub = l[0] + '_' + l[1]
        mutys = [];
        unmutys = [];

        if mother_flag == 0 and father_flag == 0:
            concluall = concluall + "，";
            conclu = u'由于先证者的父母未验证，无法确定变异的遗传来源;';
        elif mother_flag == 1 and father_flag == 1:
            print (siteGenotype[propositusID][sitesub])
            print (siteGenotype[mother_ID][sitesub])
            print (siteGenotype[father_ID][sitesub])
            if siteGenotype[propositusID][sitesub] == siteGenotype[mother_ID][sitesub] and siteGenotype[propositusID][
                sitesub] != siteGenotype[father_ID][sitesub] or siteGenotype[propositusID][sitesub] == "半合子" and \
                    siteGenotype[mother_ID][sitesub] == '变异杂合':
                concluall = concluall + "，先证者的母亲 %s 携带该变异，" % (persen['母亲'][0][2]);
                concluall = concluall + "父亲 %s 未携带该变异，" % (persen['父亲'][0][2]);
                conclu = "可见先证者的变异遗传自母亲，";
            elif siteGenotype[propositusID][sitesub] == siteGenotype[father_ID][sitesub]:
                concluall = concluall + "，先证者的父亲 %s 携带该变异，" % (persen['父亲'][0][2]);
                concluall = concluall + "母亲 %s 未携带该变异，" % (persen['母亲'][0][2]);
                conclu = "可见先证者的变异遗传自父亲，";
            else:
                concluall = concluall + "父亲 %s 、母亲 %s 未携带该变异，" % (persen['父亲'][0][2], persen['母亲'][0][2]);
                conclu = u'可见先证者的变异为denovo变异。';
        elif mother_flag == 1 and father_flag == 0:
            if siteGenotype[propositusID][sitesub] == siteGenotype[mother_ID][sitesub] or siteGenotype[propositusID][
                sitesub] == "半合子" and siteGenotype[mother_ID][sitesub] == "变异杂合":
                concluall = concluall + "，先证者的母亲 %s 携带该变异，" % (persen['母亲'][0][2]);
            else:
                concluall = concluall + "，先证者的母亲 %s 未携带该变异，" % (persen['母亲'][0][2]);
                conclu = u'由于父亲未验证，无法确定变异的遗传来源，';
        elif mother_flag == 0 and father_flag == 1:
            if siteGenotype[propositusID][sitesub] == siteGenotype[father_ID][sitesub]:
                concluall = concluall + "，先证者的父亲 %s 携带该变异，" % (persen['父亲'][0][2]);
            else:
                concluall = concluall + "，先证者的父亲 %s 未携带该变异，" % (persen['父亲'][0][2]);
                conclu = u'由于母亲未验证，无法确定变异的遗传来源，';

        for relation in family:

            if relation == "先证者" or relation == "母亲" or relation == "父亲":
                continue
            for li in persen[relation]:
                rid = li[0]
                print (rid)
                print (sitesub)
                if siteGenotype[rid][sitesub] == siteGenotype[propositusID][sitesub]:
                    mutys.append(relation + li[2])
                else:
                    unmutys.append(relation + li[2])

        if len(mutys) != 0:
            mutation = '、'.join(mutys) + "携带该变异，"
        if len(unmutys) != 0:
            unmutation = '、'.join(unmutys) + "未携带该变异，";
        concluall = concluall + mutation + unmutation + conclu + RichText('\n')

        mutation = '';
        unmutation = '';
    concluall = concluall + u'具体请结合临床分析，以下为本次检测样本的测序峰图。' + RichText('\n');
    return concluall

def set_updatefields_true(docx_path):
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    doc = Document(docx_path)
    # add child to doc.settings element
    element_updatefields = lxml.etree.SubElement(
        doc.settings.element, namespace + "updateFields"
    )
    element_updatefields.set(namespace + "val", "true")
    print (docx_path)
    docx_path2 = docx_path.replace("NN0","N0")
    print (docx_path2)
    doc.save(docx_path2)

def main_2():
    argInit()
    cover_Dic={}
    if module.strip().upper() == 'GXY':
        cover_Dic['title'] = 'Part'+part+' — 检测结果'
    else:
        cover_Dic['title'] = 'Part'+part+' — 单基因高血压/血钾异常基因检测结果'
    out_GXY=open(outDir+'/'+sampleid+'_Cardio_GXY.tex','w')
    if os.path.isfile(yinyang_file) == False:
        sys.stderr.write('%s is not a file' % yinyang_file)
        sys.exit(1)
    try:
        yinyang_handle = open(yinyang_file,'r')
        yinyang = yinyang_handle.readlines()[0].strip()
        yinyang_handle.close()
        if yinyang == '0':
            id_filter,id_filterfu = get_filter(filterFile,sampleid)
            with open('/PUBLIC/pipline/script/Report/XKFW/V2_171130/tempt/Cardio_GXYyin_temple.tex','r') as file_handle:
                GXY_temp = Template(file_handle.read())
                diseaseDic,siteDic,siteDic_fu = prepare_allDic(id_filter,disease_file,id_filterfu)
                fubiao_table = get_fubiao_table(siteDic_fu)
                cover_Dic.update(fubiao_table)
                num=len(id_filterfu)
                if num==0:
                    cover_Dic['yin']=r'''此次单基因高血压/血钾异常基因检测，包含41个疾病相关基因，筛查15种疾病亚型。检测结果显示您并未发生明确致病变异的位点，因此，您的疾病症状并不是由这些已知基因的致病变异引起的，可能是目前研究尚未涉及的基因变异，也可能是由于其他因素，如不良生活习惯、环境、其他疾病等引起。建议您注意生活细节，积极改变不良生活方式，并配合医生治疗。 

以上检测结果仅基于当前的科学研究，随研究文献的更新，疾病-基因的对应关系可能会发生变化，因此我们的报告具有一定的时效性 。
'''             
                else:
                    cover_Dic['yin']=r'''此次单基因高血压/血钾异常基因检测，包含41个疾病相关基因，筛查15种疾病亚型。检测结果显示您并未发生高风险的位点变异；一些变异位点由于相关研究较少、ACMG证据条目不足，基于现有研究判断其致病性等级较低，但不排除导致疾病的可能性，这类位点的致病性等级及相关证据条目展示在附表中，供您和临床医生参考。

若检出位点的相关疾病与您现有的临床表型不符，那么您的相关临床症状可能与其他未知基因变异、疾病的不完全外显性、不良生活习惯或环境等因素相关。建议您注意生活细节，积极改变不良生活方式，或去医院进行相关疾病的进一步检测，听从医生的临床建议，定期体检，并采取积极的生活干预。

若您未发现任何相关疾病表型，则这类 ACMG 致病类证据较少的位点可能并不致病。 

以上检测结果仅基于当前的科学研究，随研究文献的更新，位点致病性情况可能会发生变化，属于正常现象。如果您想进一步了解基因变异的信息或家族遗传史，建议联系家人一起进行检测。
'''
                #module_li=sorted([x.strip().upper() for x in module.strip().split('+')])
                #gxy_module=['GXY','DRUGB']
                #module_gxy=sorted(['GXY','DRUGB'])
                #module_gxy2=sorted(['GXY','DRUGB','CS'])
                #if module_li == module_gxy: cover_Dic['yin']=r'''
#\fontsize{10.5}{13.65}\selectfont{\color{MyFontGray} 为了使您更有效的服用降压药，我们还为您检测了 40 种心血管常用药物的敏感基因位点，涵盖降压药、降脂药、降糖药和抗凝药，并给出了具体用药指导,请您详见 Part3 部分报告。}\\ \\'''
                #elif module_li == module_gxy2: cover_Dic['yin']=r'''
#\fontsize{10.5}{13.65}\selectfont{\color{MyFontGray} 为了使您更有效的服用降压药，我们还为您检测了 40 种心血管常用药物的敏感基因位点，涵盖降压药、降脂药、降糖药和抗凝药，并给出了具体用药指导,请您详见 Part4 部分报告。}\\ \\'''
                #else: cover_Dic['yin']=''
                print (cover_Dic)
                report = GXY_temp.safe_substitute(cover_Dic)
                print >> out_GXY,report
        elif yinyang == '1':
            id_filter,id_filterfu = get_filter(filterFile,sampleid)
            with open('/PUBLIC/pipline/script/Report/XKFW/V2_171130/tempt/Cardio_GXYyang_temple.tex','r') as file_handle:
                GXY_temp = Template(file_handle.read())
                diseaseDic,siteDic,siteDic_fu = prepare_allDic(id_filter,disease_file,id_filterfu)
                #zongshu = get_zongshu(siteDic,diseaseDic)
                zongshu = get_zongshu(siteDic)
                disease_table = get_disease_table(diseaseDic,siteDic)
                site_table = get_site_table(siteDic)
                fubiao_table = get_fubiao_table(siteDic_fu)
                jiexi_table = get_jiexi_table(diseaseDic,siteDic)
                cover_Dic.update(zongshu)
                cover_Dic.update(disease_table)
                cover_Dic.update(site_table) 
                cover_Dic.update(fubiao_table)
                cover_Dic.update(jiexi_table)
                report = GXY_temp.safe_substitute(cover_Dic)
                print >> out_GXY,report
    except Exception as e:
        traceback.print_exc(e)
    finally:
        out_GXY.close()

def main():
    argInit()
    context = {}
    referencearr = []

    # tpl=DocxTemplate('南方医院-基因检测报告模版20190517.docx')
    if os.path.isfile(yinyang_file) == False:
        sys.stderr.write('%s is not a file' % yinyang_file)
        sys.exit(1)
    try:
        yinyang_handle = open(yinyang_file,'r')
        yinyang = yinyang_handle.readlines()[0].strip()
        yinyang_handle.close()
    except Exception as e:
        traceback.print_exc(e)
    finally:
        pass

    id_filter, id_filterfu = get_filter(filterFile,sampleid)
    diseaseDic,siteDic,siteDic_fu = prepare_allDic(id_filter,disease_file,id_filterfu)# wph add 190325
    fubiao_table = get_fubiao_table(siteDic_fu)
    sampleInf_Dic = get_sampleInf(sampleid, cfgFile, feature_file)
    QC_Inf_Dic = get_QC_Inf(QCfile, sampleid)
    if yinyang =="1":
        print (type(yinyang))
        disease_table = get_disease_table(diseaseDic, siteDic)
        site_table = get_site_table(siteDic)
        zongshu = get_zongshu(siteDic)
        jiexi_table = get_jiexi_table(diseaseDic, siteDic)
        print (jiexi_table)
        Sangersite = get_sanger_Inf(siteDic)
        SangerFamily = get_sanger_Inf_Family()

        context['yinyang'] = yinyang
        context['detect_summary'] = zongshu['detect_summary']
        context['tbl_contents'] = disease_table['tbl_contents']
        context['tbl_contents2'] = site_table['tbl_contents2']
        context['genetic_model'] = disease_table['genetic_model']
        context['jiexi_site'] = jiexi_table['jiexi_site']
        context['varient'] = Sangersite['varient']
        if SangerFamily:
            context['sangerSam'] = SangerFamily['sangerSam']
            context['sangerResult'] = SangerFamily['sangerResult']
            context['varientFamily'] = SangerFamily['varientFamily']
            context['col_labels'] = SangerFamily['col_labels']
            context['descript'] = SangerFamily['descript']


    context['tbl_contents3'] = fubiao_table['tbl_contents3']



    '''for item in disease_table['type_lite']:
        if item not in referencearr:
            referencearr.append(item)
    for item in jiexi_table['wenxian']:
        if item not in referencearr:
            referencearr.append(item)'''
    for item in fubiao_table['wenxian']:
        if item not in referencearr:
            referencearr.append(item)

    referencearr.append('Richards Sue,Aziz Nazneen,Bale Sherri,Bick David,Das Soma,Gastier-Foster Julie,Grody Wayne W,Hegde Madhuri,Lyon Elaine,Spector Elaine,Voelkerding Karl,Rehm Heidi L,ACMG Laboratory Quality Assurance Committee.Standards and guidelines for the interpretation of sequence variants: a joint consensus recommendation of the American College of Medical Genetics and Genomics and the Association for Molecular Pathology.[J].Genet. Med.2015 May;17(5):405-24.')

    lite = []
    for index, item in enumerate(referencearr):
        indexs = index + 1
        indexsa = str(indexs) + '.'
        refline = {'index': indexsa, 'con': item}
        lite.append(refline)
    context['reference'] = lite

    for key in QC_Inf_Dic:
        context[key] = QC_Inf_Dic[key]

    for key in sampleInf_Dic:
        context[key] = sampleInf_Dic[key]

    context['telphone'] = '-'
    # context['break']=tpl.add_page_break()
    context['program'] = '高血压基因检测Panel'
    context['testMethmod'] = '高通量测序技术'
    context['sampletype'] = '-'
    context['laboratory'] = '-'
    context['doctor'] = '-'
    context['reportHeader'] = context['name']
    context['reportNumber'] = '-'

    sd2 = tpl.new_subdoc()
    sd2.add_page_break()
    context['break'] = sd2
    tpl.render(context, autoescape=True)
    outfile = outDir +'/南方医院-诺禾心康-N' + sampleid  + '_' + context['name'] + '_单基因高血压基因检测报告.docx'
    tpl.save(outfile)
    set_updatefields_true(outfile)

if __name__ == '__main__':
    main()
